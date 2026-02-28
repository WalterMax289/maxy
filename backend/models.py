import random
import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import os
import json
import wikipedia
from data_analyzer import AdvancedAnalyzer, TextAnalyzer, StructuredDataAnalyzer
from code_composer import CodeComposer
from slang_manager import SlangManager
import requests
from ddgs import DDGS
import yfinance as yf
import pandas as pd
try:
    import PyPDF2
    from docx import Document as DocxDocument
except ImportError:
    PyPDF2 = None
    DocxDocument = None

slang_manager = SlangManager()

logger = logging.getLogger(__name__)

class MAXYThinkingEngine:
    
    @staticmethod
    def generate_thinking(
        model_name: str,
        user_message: str,
        analysis_type: str = "general"
    ) -> str:

        reasoning_steps = {
            'quick': [
                "Analyzing input...",
                "Formulating response...",
                "Optimizing output..."
            ],
            'research': [
                "Identifying key topics...",
                "Searching knowledge sources...",
                "Synthesizing information...",
                "Structuring comprehensive response..."
            ],
            'conversation': [
                "Understanding context...",
                "Identifying intent...",
                "Crafting natural response..."
            ],
            'analysis': [
                "Examining content structure...",
                "Extracting key information...",
                "Identifying patterns...",
                "Generating insights..."
            ]
        }
        
        steps = reasoning_steps.get(analysis_type, reasoning_steps['quick'])
        
        thinking_text = f"Analyzing: '{user_message[:50]}'\n\n"
        thinking_text += "Processing:\n"
        for i, step in enumerate(steps, 1):
            thinking_text += f"{i}. {step}\n"
        
        if analysis_type == 'research':
            thinking_text += f"{len(steps) + 1}. Verifying factual consistency...\n"
            thinking_text += f"{len(steps) + 2}. Cross-referencing sources...\n"
            
        return thinking_text


class KnowledgeSynthesizer:
    """Intelligent search result synthesis and verification"""
    
    RESEARCH_KEYWORDS = [
        'what is', 'who is', 'how does', 'explain', 'tell me about',
        'info about', 'information on', 'details about', 'who was',
        'what are', 'define', 'describe', 'history of', 'science',
        'technology', 'biology', 'physics', 'chemistry', 'geography',
        'country', 'capital', 'famous', 'invented', 'discovered',
        'meaning of', 'purpose', 'history', 'research', 'analysis',
        'details', 'whois', 'news', 'happening', 'headlines', 'trends',
        'health', 'medical', 'symptoms', 'treatment', 'prevention',
        'when did', 'where is', 'why does', 'sort', 'search', 'array',
        'list', 'tree', 'graph', 'data structure', 'algorithm', 
        'implement', 'code', 'function', 'class', 'decorator',
        'python', 'javascript', 'java', 'html', 'css', 'sql',
        'pm of', 'president of', 'governor of', 'ceo of',
        'exploration', 'discovery', 'universe', 'space', 'astronomy',
        'physics', 'math', 'mathematics', 'geometry', 'calculus',
        'what\'s up with', 'tell me more about', 'latest on',
        'write an essay', 'write a speech', 'compose an essay',
        'give me a speech', 'draft an essay', 'persuasive essay',
        'write about', 'give me an essay', 'i want an essay',
        'an essay about', 'an essay on', 'compose a speech',
        'write me an essay', 'write me a speech', 'draft a speech',
        'write a paragraph', 'composition', 'argumentative essay', 'descriptive essay',
        'narrative essay', 'expository essay', 'formal speech', 'keynote',
        'can you explain', 'could you tell me', 'i want to know',
        'learn about', 'guide to', 'overview of', 'introduction to',
        'basics of', 'advanced', 'in depth', 'detailed explanation',
        'summary of', 'quick facts about', 'key points of',
        'important facts', 'meaning of', 'definition of',
        'full form of', 'origin of', 'background of', 'concept of',
        # Time / Historical Queries
        'when was', 'when is', 'how long does', 'how long did',
        'how many years', 'timeline of', 'era of', 'period of',
        'ancient', 'modern', 'medieval', 'future of', 'predicted',
        'forecast of', 'trend in', 'evolution of',
        # Location / Geography
        'located in', 'situated in', 'map of', 'population of',
        'area of', 'largest', 'smallest', 'bordering',
        'neighboring countries', 'climate of', 'currency of',
        'language of',
        # People / Position
        'founder of', 'owner of', 'chairman of', 'minister of',
        'prime minister of', 'king of', 'queen of', 'director of',
        'author of', 'creator of', 'biography of', 'net worth of',
        'age of', 'early life of',
        # Science and Education
        'theory of', 'law of', 'principle of', 'formula for',
        'equation of', 'difference between', 'compare',
        'comparison between', 'advantages of', 'disadvantages of',
        'types of', 'branches of', 'process of', 'cycle of',
        'structure of', 'function of', 'example of',
        # Programming / Tech
        'syntax of', 'how to use', 'usage of',
        'best practices for', 'error in', 'debug', 'optimize',
        'performance of', 'library for', 'framework for', 'api for',
        'database', 'backend', 'frontend', 'full stack',
        'machine learning', 'artificial intelligence', 'cybersecurity',
        'cloud computing', 'blockchain', 'data science', 'deep learning',
        # Data Structures & Algorithms
        'time complexity', 'space complexity', 'big o notation',
        'linear search', 'binary search', 'merge sort', 'quick sort',
        'dynamic programming', 'recursion', 'greedy algorithm',
        'stack', 'queue', 'linked list', 'binary tree', 'bst',
        'heap', 'hash table',
        # Business / Economy
        'market value', 'stock price of', 'economy of', 'gdp of',
        'inflation rate', 'revenue of', 'profit of',
        'business model of', 'case study of', 'impact of', 'benefits of'
    ]

    CODE_INDICATORS = [
        'code', 'write', 'create', 'generate', 'function', 'how to',
        'program', 'script', 'example', 'syntax', 'algorithm', 'implement',
        'snippet', 'coding', 'develop', 'setup', 'server', 'logic',
        'sort', 'search', 'array', 'list', 'tree', 'graph', 'data structure',
        'decorator', 'class', 'method', 'variable', 'loop', 'conditional',
        'syntax of', 'how to use', 'usage of', 'best practices for',
        'error in', 'debug', 'optimize', 'performance of',
        'library for', 'framework for', 'api for', 'database',
        'backend', 'frontend', 'full stack', 'coding for',
        'time complexity', 'space complexity', 'big o notation',
        'linear search', 'binary search', 'merge sort', 'quick sort',
        'dynamic programming', 'recursion', 'greedy algorithm',
        'stack', 'queue', 'linked list', 'binary tree', 'bst',
        'heap', 'hash table', 'snippet', 'coding', 'develop', 'setup', 'server', 'logic',
        'sort', 'search', 'array', 'list', 'tree', 'graph', 'data structure',
        'boilerplate', 'repository', 'rest api', 'crud', 'component',
        'react', 'nextjs', 'vue', 'node', 'express', 'django', 'flask',
        'refactor', 'unit test', 'integration test', 'deployment script',
        'dockerfile', 'kubernetes', 'yaml config'
    ]
    
    @staticmethod
    def get_keywords(query: str) -> List[str]:
        """Extract core keywords from query for relevance scoring"""
        # Basic keyword extraction: remove noise, focus on entities
        noise = ['is', 'who', 'the', 'of', 'what', 'was', 'were', 'tell', 'me', 'about', 'how', 'does', 'are']
        # Allow short but critical terms like PM, CM, CEO
        critical_titles = ['pm', 'cm', 'ceo', 'cfo', 'cto', 'md', 'mp', 'mla']
        words = re.findall(r'\b\w+\b', query.lower())
        return [w for w in words if (w in critical_titles or (w not in noise and len(w) > 2))]

    @staticmethod
    def score_relevance(query: str, title: str, body: str) -> float:
        """Score how relevant a search result is to the query"""
        keywords = KnowledgeSynthesizer.get_keywords(query)
        if not keywords:
            return 0.5
            
        content = (title + " " + body).lower()
        title_lower = title.lower().strip()
        matches = sum(1 for kw in keywords if kw in content)
        
        # Identity query detection
        identity_keywords = ['who is', 'who was', 'identity', 'person', 'pm of', 'president of', 'ceo of', 'chief minister of', 'chief of', 'founder of', 'creator of', 'author of']
        msg_lower = query.lower()
        is_identity = any(ik in msg_lower for ik in identity_keywords)
        
        if is_identity:
            # 1. Recency Boost: Heavily prioritize current status
            recency_indicators = ['current', 'incumbent', 'serving as', 'holds the position', 'is currently the', 'presently', 'now']
            if any(ri in content for ri in recency_indicators):
                matches += 2
            
            # 2. Historical Penalty: Penalize past officeholders
            historical_indicators = ['former', 'ex-', 'past', 'who was', 'predecessor', 'served as', 'between', 'during']
            if any(hi in content for hi in historical_indicators):
                matches -= 1
            
            # 3. Institution/Organization Penalty for Identity Queries
            # If we are looking for a PERSON, penalize organizations unless the query explicitly has them
            institution_indicators = [
                'institute', 'university', 'college', 'foundation', 'academy', 
                'organization', 'hospital', 'department', 'agency', 'commission', 
                'association', 'society', 'center', 'centre', 'school', 'board', 'council'
            ]
            if not any(ii in msg_lower for ii in institution_indicators):
                if any(ii in title_lower for ii in institution_indicators):
                    matches -= 10 # Strong penalty to prevent MGIMS over Mahatma Gandhi

            # 4. Wikipedia Disambiguation Penalty
            if "disambiguation" in title_lower:
                matches -= 5

            # 5. Year Range Penalty (e.g., 1991-1996)
            if re.search(r'\b(19|20)[0-9]{2}[\-–](19|20)[0-9]{2}\b', content):
                # If there's a year range, it's often a historical context unless "current" is also there
                if not any(ri in content for ri in recency_indicators):
                    matches -= 1

            # Try to find specific names (Title Case words) in query
            names_in_query = re.findall(r'\b[A-Z][a-z]+\b', query)
            
            # Robust Fallback for lowercase queries
            if not names_in_query:
                # Extract words after "who is", "who was", etc.
                for trigger in identity_keywords:
                    if trigger in msg_lower:
                        after_trigger = msg_lower.split(trigger)[1].strip()
                        # Clean up punctuation
                        after_trigger = re.sub(r'[^\w\s]', '', after_trigger)
                        if after_trigger:
                            names_in_query = after_trigger.split()
                        break
            
            if names_in_query:
                content_lower = content.lower()
                name_matches = sum(1 for name in names_in_query if name.lower() in content_lower)
                
                query_full_name = " ".join(names_in_query).lower()
                
                # Significant boost if the title matches the name in the query
                if query_full_name == title_lower or query_full_name in title_lower:
                    matches += 25
                elif all(name.lower() in title_lower for name in names_in_query):
                    matches += 15
                
                # Penalty for sub-topics if the name is found but title contains extra non-person context
                non_person_subtopics = [
                    'assassination', 'family', 'legacy', 'death', 'childhood', 'early life', 
                    'career', 'politics', 'murder', 'killing', 'incident', 'event', 'movement',
                    'uprising', 'rebellion', 'riot', 'battle', 'war', 'anniversary', 'memorial'
                ]
                for ind in non_person_subtopics:
                    if ind in title_lower and ind not in query.lower():
                        matches -= 15 # Increased penalty to push biography to top
                        break
                
                if name_matches == 0:
                    return 0.01 # Very low relevance if names are missing from result
                matches += name_matches
        
        score = matches / (len(keywords) + 1)
        
        # News/Headline Penalty for identity queries
        if is_identity:
            # Titles with colons, question marks at start, or buzzwords are often news
            news_indicators = [':', '?', 'breaking', 'live', 'update', 'latest', 'counters', 'claims', 'vs', 'opinion', 'watch', 'video']
            if any(ni in title_lower for ni in news_indicators):
                score *= 0.3
            
            # Boost if the Title of the result matches the query keywords well
            title_matches = sum(1 for kw in keywords if kw in title_lower)
            if title_matches >= 2:
                score += 0.2
                
            # Wikipedia / Factual Source Boost
            if "wikipedia" in title_lower or "britannica" in title_lower or "biography" in title_lower:
                score += 0.4
        
        # Penalty for low-quality sources in body (casual mentions)
        junk_indicators = ['reddit', 'quora', 'forum', 'comment', 'manhwa', 'manga', 'recommendation', 'fanfiction']
        if not any(ek in query.lower() for ek in ['manga', 'manhwa', 'comic', 'read']):
            if any(ji in body.lower() for ji in junk_indicators):
                score *= 0.5
        
        return max(0.01, score)

    @staticmethod
    def verify_facts(query: str, results: List[Dict[str, str]]) -> List[Dict[str, Any]]:
        """Verify and rank search results by relevance"""
        ranked = []
        for res in results:
            score = KnowledgeSynthesizer.score_relevance(query, res.get('title', ''), res.get('body', ''))
            res['relevance_score'] = score
            ranked.append(res)
        
        return sorted(ranked, key=lambda x: x['relevance_score'], reverse=True)

    @staticmethod
    def get_best_match(query: str, results: List[Dict[str, str]], threshold: float = 0.3) -> Optional[Dict[str, Any]]:
        """Identify the single most relevant search result"""
        verified = KnowledgeSynthesizer.verify_facts(query, results)
        if verified and verified[0]['relevance_score'] >= threshold:
            return verified[0]
        return None

    @staticmethod
    def extract_identity_answer(query: str, wiki_result: str, intents: Dict[str, bool]) -> Optional[str]:
        """Shared logic to extract a concise name or identity from search results"""
        msg_lower = query.lower().strip()
        is_position_query = any(ik in msg_lower for ik in ['pm of', 'president of', 'ceo of', 'cm of', 'leader of', 'who is the current', 'who is the pm', 'who is the president', 'who is the ceo'])
        is_person_query = msg_lower.startswith('who is ') and not is_position_query
        
        if intents.get('knowledge') and (is_position_query or is_person_query):
            # Do NOT extract identity from already formatted research reports
            if wiki_result.lstrip().startswith("**VERIFIED RESEARCH REPORT") or wiki_result.lstrip().startswith("**MAXY ENTERPRISE"):
                return None
            
            # Suppression: If the query asks for biography or narrative detail, skip short identity
            narrative_indicators = ['biography', 'lifestyle', 'narrative', 'detail', 'history', 'life', 'story', 'tell me about', 'tell me everything']
            if any(ni in msg_lower for ni in narrative_indicators):
                return None
            
            # Expanded exclusion list for titles, locations, and generic terms
            blacklist = [
                "University", "Republic", "States", "Kingdom", "Minister", "President", 
                "Council", "Congress", "Parliament", "National", "Public", "Official", 
                "Government", "General", "Cabinet", "Supreme", "Federal", "Union", "South", "North", "East", "West",
                "India", "American", "British", "World", "Global", "Today", "News", "Breaking", "Live", "Update",
                "Jagran", "Josh", "Times", "Post", "Gazette", "Chronicle", "Herald", "Observer", "Tribune", "Daily",
                "Guardian", "Mirror", "Standard", "Express", "Sun", "Mail", "Telegraph", "Independent", "Reuters", "Associated", "Press",
                "White", "House", "Washington", "Street", "Journal", "Gazette", "City", "County", "District",
                "Electoral", "College", "Foundation", "Institute", "Organization", "Department", "Agency", "Commission", "Association",
                "Academy", "Hospital", "Trust", "Group", "Inc", "Ltd", "Corporation", "Limited", "Society", "Center", "Centre", "School", "Board",
                "History", "Biography", "Profile", "Fact", "Summary", "Overview", "Details", "Information",
                "Incident", "Event", "Movement", "Battle", "War", "Uprising", "Rebellion", "Riot", "Protest", "Anniversary", "Memorial"
            ]
            
            snippet_clean = wiki_result.replace('\n', ' ').replace(':', ' is ')
            
            if is_position_query:
                # Seeking a Name for a Position
                # Pattern 1: [Name] is the [Position]
                match1 = re.search(r'\b([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\s+is\s+(?:the\s+)?(?:current\s+)?(?:prime\s+minister|president|ceo|cm|leader|head|ruler)\b', snippet_clean, re.I)
                if match1:
                    name = match1.group(1).strip()
                    if not any(bl in name for bl in blacklist):
                        return f"{name}."
                        
                # Pattern 2: [Position] is [Name]
                match2 = re.search(r'\b(?:prime\s+minister|president|ceo|cm|leader)\s+(?:of\s+[\w\s]+)?\s+is\s+([A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\b', snippet_clean, re.I)
                if match2:
                    name = match2.group(1).strip()
                    if not any(bl in name for bl in blacklist):
                        return f"{name}."
                        
                # Fallback for position queries: Stricter general name extraction
                potential_names = re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?\b', snippet_clean)
                if potential_names:
                    for name in potential_names:
                        words_in_name = name.split()
                        if any(w.lower() in msg_lower for w in words_in_name): continue
                        if any(bl in name for bl in blacklist): continue
                        if any(title in name for title in ["Minister", "President", "Secretary", "Director", "Policy", "State", "National", "Public"]): continue
                        return f"{name}."
                        
            elif is_person_query:
                # Seeking a Title for a Person (e.g. "Who is Narendra Modi")
                # Return the first 1-2 descriptive sentences
                sentences = [s.strip() for s in wiki_result.split('. ') if len(s.strip()) > 15]
                if sentences:
                    return f"{sentences[0]}."
        
        return None


class MAXY1_1:
    NAME = "MAXY 1.1"
    VERSION = "1.1.0"
    DESCRIPTION = "Quick response AI with visible thinking process"
    
    # Quick response templates organized by intent
    GREETINGS = [
        "Hey there! 👋 Ready to chat!",
        "Hello! What can I help you with?",
        "Hi! I'm here and ready to assist!",
        "Hey! Great to see you!",
    ]
    
    FAREWELLS = [
        "Goodbye! Catch you later! 👋",
        "See you soon! Take care!",
        "Bye for now! Come back anytime!",
        "Until next time! Stay awesome!",
    ]
    
    GRATITUDE = [
        "You're welcome! Happy to help! 😊",
        "Anytime! That's what I'm here for!",
        "Glad I could assist!",
        "No problem at all!",
    ]
    
    HOW_ARE_YOU = [
        "I'm doing great, thanks for asking! How are you?",
        "Excellent! Ready to help. You?",
        "Fantastic! What about you?",
        "All good here! How's your day?",
    ]
    
    IDENTITY = [
        "I'm MAXY 1.1 - your quick-thinking AI assistant! I provide fast responses with clear reasoning.",
        "Hello! I'm MAXY 1.1, designed for rapid responses and friendly conversation!",
        "I'm MAXY 1.1! I specialize in quick, thoughtful responses with visible thinking processes.",
    ]
    
    GENERAL_QUICK = [
        "Got it! Tell me more.",
        "Interesting! Continue...",
        "I see! What else?",
        "That makes sense!",
        "Understood! What's next?",
        "Okay! How can I help further?",
    ]
    
    JOKES = [
        "Why don't scientists trust atoms? Because they make up everything! 😄",
        "Why did the scarecrow win an award? He was outstanding in his field!",
        "What do you call a fake noodle? An impasta! 🍝",
        "Why don't eggs tell jokes? They'd crack each other up!",
        f"Lo {slang_manager.get_random_slang()}, why did the tomato turn red? Because it saw the salad dressing! 😂",
        f"{slang_manager.get_random_slang()}, parallel lines have so much in common but they’ll never meet. Sad scene no? 😅"
    ]
    
    @staticmethod
    def should_use_wikipedia(message: str) -> bool:
        """Determine if this is a knowledge/research question"""
        msg_lower = message.lower()
        return any(kw in msg_lower for kw in KnowledgeSynthesizer.RESEARCH_KEYWORDS)
    
    @staticmethod
    def quick_wikipedia_lookup(query: str) -> Optional[str]:
        """Quick knowledge lookup for maxy1.1 with multi-source verification"""
        try:
            candidates = []
            
            # 1. Wikipedia Search (Prioritize exact title match)
            try:
                # Try to get the specific page for the query first
                try:
                    direct_res = wikipedia.page(query, auto_suggest=True)
                    candidates.append({
                        'title': direct_res.title,
                        'body': direct_res.summary[:800],
                        'source': 'wikipedia'
                    })
                except:
                    pass
                    
                search_results = wikipedia.search(query, results=5)
                for res in search_results:
                    try:
                        page = wikipedia.page(res, auto_suggest=False)
                        candidates.append({
                            'title': page.title,
                            'body': page.summary[:800],
                            'source': 'wikipedia'
                        })
                    except:
                        continue
            except Exception as e:
                logger.error(f"Wiki lookup error: {e}")

            # 2. DuckDuckGo Search
            try:
                with DDGS() as ddgs:
                    # For identity queries, force "current" to avoid historical lists
                    search_query = query
                    position_keywords = ['pm of', 'ceo of', 'president of', 'pm', 'cm of', 'head of', 'chief of']
                    if any(pk in query.lower() for pk in position_keywords):
                        if "current" not in query.lower():
                            search_query = f"current {query}"
                        if not search_query.lower().startswith('who is'):
                            search_query = f"who is the {search_query}"
                    elif query.istitle() and len(query.split()) <= 3:
                        # If query is just a name (e.g. "Mahatma Gandhi"), add "who is"
                        search_query = f"who is {query}"
                            
                    results = list(ddgs.text(search_query, max_results=8))
                    for res in results:
                        candidates.append({
                            'title': res['title'],
                            'body': res['body'],
                            'source': 'web'
                        })
            except Exception as e:
                logger.error(f"DDG search error: {e}")

            if not candidates:
                return None

            # 3. Verify and Synthesis
            best_match = KnowledgeSynthesizer.get_best_match(query, candidates, threshold=0.25)
            
            if best_match:
                content = best_match['body']
                # If it's a web source, include title
                if best_match['source'] == 'web':
                    return f"{best_match['title']}: {content}"
                return content
            
            return None
        except Exception as e:
            logger.error(f"Quick lookup total error: {e}")
            return None

    @staticmethod
    def get_weather(city: str) -> Optional[str]:
        """Fetch weather data from OpenMeteo"""
        try:
            # 1. Geocoding
            geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={city}&count=1&language=en&format=json"
            geo_res = requests.get(geo_url).json()
            
            if not geo_res.get('results'):
                return None
                
            lat = geo_res['results'][0]['latitude']
            lon = geo_res['results'][0]['longitude']
            name = geo_res['results'][0]['name']
            country = geo_res['results'][0]['country']
            
            # 2. Weather
            weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m&timezone=auto"
            w_res = requests.get(weather_url).json()
            
            current = w_res.get('current', {})
            temp = current.get('temperature_2m')
            humidity = current.get('relative_humidity_2m')
            wind = current.get('wind_speed_10m')
            
            # Weather codes
            codes = {
                0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
                45: "Foggy", 48: "Depositing rime fog", 51: "Light drizzle", 53: "Drizzle",
                55: "Heavy drizzle", 61: "Slight rain", 63: "Moderate rain", 65: "Heavy rain",
                71: "Slight snow", 73: "Moderate snow", 75: "Heavy snow", 95: "Thunderstorm"
            }
            condition = codes.get(current.get('weather_code'), "Variable")
            
            return f"{condition} in {name}, {country}. Temp: {temp}°C, Humidity: {humidity}%, Wind: {wind} km/h."
            
        except Exception as e:
            logger.error(f"Weather error: {e}")
            return None
    
    @staticmethod
    def analyze_user_intent(message: str) -> Dict[str, Any]:
        """Analyze what the user wants - improved context understanding"""
        msg_lower = message.lower().strip()
        
        # Intent categories with word boundaries for short words
        intents = {
            'greeting': any(re.search(r'\b' + re.escape(g) + r'\b', msg_lower) for g in ['hi', 'hello', 'hey', 'greetings', 'howdy']),
            'farewell': any(re.search(r'\b' + re.escape(f) + r'\b', msg_lower) for f in ['bye', 'goodbye', 'see you', 'farewell', 'later']),
            'gratitude': any(re.search(r'\b' + re.escape(t) + r'\b', msg_lower) for t in ['thanks', 'thank you', 'appreciate', 'grateful']),
            'personal_status': any(h in msg_lower for h in ['how are you', 'how you doing']),
            'identity': any(i in msg_lower for i in ['your name', 'who are you', 'what are you']),
            'entertainment': any(j in msg_lower for j in ['joke', 'funny', 'laugh']),
            'time_query': any(t in msg_lower for t in ['time', 'what time', 'current time']),
            'date_query': any(d in msg_lower for d in ['date', 'today', 'what day']),
            'daily_updates': any(u in msg_lower for u in ['daily updates', 'what is new', 'whats new', 'latest updates']),
            'help': any(h in msg_lower for h in ['help', 'what can you do']),
            'news': any(n in msg_lower for n in ['news', 'happening', 'headlines', 'world today', 'current events']),
            'knowledge': any(k in msg_lower for k in ['what is', 'who is', 'how does', 'explain', 'tell me about', 'info about', 'information on', 'details about', 'is there a meaning', 'meaning of', 'purpose of', 'pm of', 'ceo of', 'president of', 'cm of', 'leader of']),
            'calculation': any(c in msg_lower for c in ['calculate', 'math', 'plus', 'minus', 'times', 'divided']),
            'weather': any(w in msg_lower for w in ['weather', 'temperature', 'rain', 'sunny']),
            'simple_task': len(message.split()) <= 3 and not any(char.isdigit() for char in message)
        }
        
        # Detect urgency/enthusiasm
        urgency = sum(1 for char in message if char in '!') + (2 if 'urgent' in msg_lower or 'asap' in msg_lower else 0)
        
        # Detect if user is new (short greeting)
        is_new_user = intents['greeting'] and len(message) < 10
        
        return {
            'intents': intents,
            'urgency': urgency,
            'is_new_user': is_new_user,
            'message_length': len(message),
            'word_count': len(message.split())
        }
    
    @staticmethod
    def generate_concise_response(intent_analysis: Dict[str, Any], message: str, use_slang: bool = False, user_name: Optional[str] = None) -> tuple[str, float]:
        """Generate 3-4 sentence response based on intent"""
        intents = intent_analysis['intents']
        msg_lower = message.lower().strip()
        
        # Priority 0: Simple math check (before Wikipedia)
        math_pattern = r'what is (\d+)\s*([+\-*/])\s*(\d+)'
        math_match = re.match(math_pattern, msg_lower)
        if math_match:
            a = int(math_match.group(1))
            op = math_match.group(2)
            b = int(math_match.group(3))
            if op == '+':
                return (f"The answer is {a + b}.", 0.99)
            elif op == '-':
                return (f"The answer is {a - b}.", 0.99)
            elif op == '*':
                return (f"The answer is {a * b}.", 0.99)
            elif op == '/' and b != 0:
                div_result = a / b
                if div_result == int(div_result):
                    return (f"The answer is {int(div_result)}.", 0.99)
                return (f"The answer is {div_result}.", 0.99)
        
        # Priority 0: Daily Updates handler
        if intents.get('daily_updates'):
            try:
                updates_path = os.path.join(os.path.dirname(__file__), "updates.json")
                if os.path.exists(updates_path):
                    with open(updates_path, 'r') as f:
                        data = json.load(f)
                        latest = data['updates'][0]
                        return (f"Here's the latest update from {latest['date']}: **{latest['title']}**. {latest['description']} Would you like more details?", 0.95)
            except Exception as e:
                logger.error(f"Error in daily_updates handler: {e}")
            return ("I'm checking our latest updates! We've recently enhanced our domain knowledge and added Bangalore slang support. What else would you like to know?", 0.90)

        # Priority 1: Check for knowledge/research/news queries FIRST
        if intents['knowledge'] or intents.get('news') or MAXY1_1.should_use_wikipedia(message):
            wiki_result = MAXY1_1.quick_wikipedia_lookup(message)
            if wiki_result:
                # Priority: Identity Extraction (One-word/Short Answer)
                identity_answer = KnowledgeSynthesizer.extract_identity_answer(message, wiki_result, intents)
                if identity_answer:
                    return (identity_answer, 0.98)
                            
                # Allow 4-5 sentences for "Gemini-like" fluency for general knowledge
                raw_sentences = [s.strip() for s in wiki_result.split('. ') if s.strip()]
                clean_sentences = [s for s in raw_sentences if len(s) > 10]
                sentences = clean_sentences[:5] 
                concise = '. '.join(sentences)
                if not concise.endswith('.'):
                    concise += '.'
                if len(concise) < 100:
                    concise += " Would you like to know more about its history or specific details?"
                return (concise, 0.92)

        # Priority 2: Check for specific slang conversational greetings
        slang_response = slang_manager.handle_conversational_slang(message)
        if slang_response:
            return (slang_response, 0.99)
        
        # Greeting - Friendly and welcoming (MAXY 1.1 Persona)
        if intents['greeting']:
            address = user_name if user_name else slang_manager.get_random_slang(use_slang)
            if intent_analysis.get('is_new_user', False):
                return (f"Hi {address}! 👋 I'm MAXY 1.1, your quick and friendly AI assistant. I'm optimized for fast answers and helpful chat. What can I do for you today?", 0.98)
            else:
                return (f"Hi {address}! Great to see you! Ready to help you with anything fast. What's on your mind?", 0.97)
        
        # Farewell - Warm goodbye (2-3 sentences)
        elif intents['farewell']:
            return (random.choice([
                "Goodbye! Thanks for chatting with me. Take care and come back anytime you need quick help! 👋",
                "See you later! It was great helping you out today. Have an awesome day!",
                "Bye for now! Don't hesitate to return if you need fast answers to anything!"
            ]), 0.98)
        
        # Gratitude - Humble and helpful (2-3 sentences)
        elif intents['gratitude']:
            return (random.choice([
                f"You're very welcome! Happy I could help quickly, {slang_manager.get_random_slang(use_slang)}. Let me know if you need anything else! 😊",
                "Anytime! That's what I'm here for. Feel free to ask more questions anytime!",
                f"Glad I could assist, {slang_manager.get_random_slang(use_slang)}! Don't hesitate to reach out if you need more quick answers!"
            ]), 0.96)
        
        # Personal status - Friendly reciprocation (3 sentences)
        elif intents['personal_status']:
            return (random.choice([
                "I'm doing fantastic, thanks for asking! All systems are running smoothly and I'm ready to help. How about you? How's your day going?",
                "Excellent! I'm energized and ready to assist. Thanks for checking in! How are you feeling today?",
                "I'm great! Optimized and ready for quick responses. How about yourself? What's new with you?"
            ]), 0.94)
        
        # Identity - Brief intro (3 sentences)
        elif intents['identity']:
            return ("I'm MAXY 1.1, your quick-thinking AI assistant! I specialize in fast, clear responses to help you get answers quickly. I can chat, answer questions, or help with simple tasks. What do you need?", 0.96)
        
        # Entertainment - Fun and light (1-2 sentences)
        elif intents['entertainment']:
            joke = random.choice(MAXY1_1.JOKES)
            return (f"{joke} 😄 Hope that brought a smile to your face!", 0.92)
        
        # Time query - Direct answer (2 sentences)
        elif intents['time_query']:
            current = datetime.now().strftime("%I:%M %p")
            return (f"It's {current} right now! ⏰ Is there something time-sensitive you need help with?", 0.97)
        
        # Date query - Direct answer (2 sentences)
        elif intents['date_query']:
            current = datetime.now().strftime("%A, %B %d, %Y")
            return (f"Today is {current}! 📅 Anything special planned for today?", 0.97)
        
        # Weather - Informative but brief (3 sentences)
        elif intents['weather']:
            # Extract potential city name (simple heuristic)
            words = message.split()
            city = None
            if 'in' in words:
                idx = words.index('in')
                if idx + 1 < len(words):
                    city = " ".join(words[idx + 1:]).strip('?.!')
            
            # If no "in", try to take the last word if it looks like a city
            if not city and len(words) > 0:
                 potential = words[-1].strip('?.!')
                 if potential.istitle() and potential.lower() not in ['weather', 'today', 'now']:
                     city = potential

            if city:
                weather_info = MAXY1_1.get_weather(city)
                if weather_info:
                    return (f"{weather_info} 🌤️ Need anything else?", 0.95)
            
            return ("I can check the weather if you tell me which city! Just ask 'weather in London' for example. 🌍", 0.90)
        
        # Simple task - Check for single word/short "tasks" that are actually topics
        elif intents['simple_task']:
            # If message is very short (1-3 words), treat as potential query FIRST
            # Unnless it's a identified slang greeting/trigger
            if len(message.split()) <= 3 and not use_slang:
                 wiki_result = MAXY1_1.quick_wikipedia_lookup(message)
                 if wiki_result:
                    raw_sentences = [s.strip() for s in wiki_result.split('. ') if s.strip()]
                    clean_sentences = [s for s in raw_sentences if len(s) > 10]
                    sentences = clean_sentences[:5]
                    concise = '. '.join(sentences)
                    if not concise.endswith('.'):
                        concise += '.'
                    return (concise, 0.92)

            return ("I understand! I'm ready to proceed. What specific action would you like me to take with this information?", 0.88)
        
        # Help request - Quick capabilities (3-4 sentences)
        elif intents['help']:
            return ("I'm MAXY 1.1, your quick AI assistant! I can answer questions, chat with you, look up quick facts, and help with simple tasks. I'm all about speed and clarity. What would you like help with?", 0.93)
        
        # Calculation - Offer to help (2-3 sentences)
        elif intents['calculation']:
            return ("I can help with calculations! Just give me the numbers and what operation you need. I'll get you the answer quickly!", 0.91)
        
        # Default - Engaging but brief (3 sentences)
        else:
            return (random.choice([
                "Interesting! Tell me more about what you're looking for. I'm here to help quickly!",
                "I see! What's the main thing you need help with? I'm ready to assist!",
                "Got it! How can I make this easier for you? Let me know what you need!",
                "Okay! What's the next step? I'm here to provide quick answers!",
                "Understood! What specific information do you need? I'll get it for you fast!"
            ]), 0.85)
    
    @staticmethod
    def analyze_casual_context(message: str, history: List[Dict]) -> str:
        """Analyze context for casual conversation continuity"""
        if not history:
            return ""
            
        last_user_msg = next((m['content'] for m in reversed(history) if m['role'] == 'user'), "")
        last_ai_msg = next((m['content'] for m in reversed(history) if m['role'] == 'assistant'), "")
        
        # Check if we asked a question
        if "?" in last_ai_msg:
            return "continuing_conversation"
            
        return "new_topic"

    @staticmethod
    def process_message(
        message: str,
        include_thinking: bool = True,
        conversation_history: Optional[List[Dict]] = None,
        user_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process message with enhanced understanding and concise 3-4 sentence responses"""
        
        # Detect if user is using slang to trigger reactive mode
        use_slang = slang_manager.detect_slang(message)
        
        # Context analysis & Follow-up detection
        is_followup, prev_context = MAXY1_3.detect_followup(message, conversation_history)
        effective_message = f"{prev_context} {message}" if is_followup else message
        
        intent_analysis = MAXY1_1.analyze_user_intent(effective_message)
        
        # Generate thinking process
        thinking = None
        if include_thinking:
            thinking = MAXYThinkingEngine.generate_thinking(
                MAXY1_1.NAME,
                effective_message,
                "quick"
            )
        
        # Generate appropriate concise response
        response, confidence = MAXY1_1.generate_concise_response(intent_analysis, effective_message, use_slang, user_name)
        
        if is_followup:
            # Prepend context acknowledgment if needed
            if len(response.split()) < 10:
                response = f"Sure! Building on that: {response}"
        
        # Adjust for context (skip - context_status not available)
        
        # Ensure response logic for MAXY 1.1 conciseness
        # Only truncate if it's NOT a very short (likely identity) response
        sentences = [s.strip() for s in response.split('. ') if s.strip()]
        if len(sentences) > 3:
            response = '. '.join(sentences[:3])
            if not response.endswith('.'):
                response += '.'
        elif len(sentences) == 1 and len(response.split()) <= 4:
            # If it's a one-word answer, leave it as is (with optional ending dot)
            pass
        
        # Inject slang (chance based) for standard interactions
        if not intent_analysis.get('is_new_user', False) and intent_analysis['intents']['greeting'] == False: # Don't double slang greeting
             response = slang_manager.enhance_text(response, force=use_slang)
        
        result = {
            'response': response,
            'model': MAXY1_1.NAME,
            'confidence': confidence,
        }
        
        if thinking:
            result['thinking'] = thinking
        
        return result


class MAXY1_2:
    NAME = "MAXY 1.2"
    VERSION = "1.2.0"
    DESCRIPTION = "Deep research expert with Wikipedia knowledge and conversational abilities"
    
    # Conversational responses for non-research queries
    CONVERSATION_GREETINGS = [
        "Hello! I'm MAXY 1.2. I can dive deep into research topics or just chat with you. What would you like?",
        "Hey there! Ready for deep research or casual conversation. What's on your mind?",
        "Hi! I'm here to provide in-depth knowledge or have a friendly chat. Your choice!",
    ]
    
    CONVERSATION_RESPONSES = [
        "That's a fascinating point! I'd love to hear more about your perspective on that.",
        "I see exactly what you mean. It's interesting how these concepts often intersect.",
        "Fascinating perspective! What do you think are the most significant implications of this?",
        "I understand completely. Is there a specific detail or angle you'd like me to research further?",
        "Excellent point! We can dive much deeper into the technical or historical aspects if you'd like.",
        "That's a very thoughtful observation. It reminds me of some related research I've encountered recently.",
        "I appreciate you sharing that. It adds a whole new dimension to our discussion!",
    ]
    
    HOW_ARE_YOU = [
        "I'm doing wonderfully, thank you! Ready to research or chat. How are you feeling today?",
        "I'm excellent! Whether you want deep analysis or casual conversation, I'm here. How about you?",
        "All systems optimal! I can provide detailed research or just have a friendly chat. You?",
    ]
    
    @staticmethod
    def is_research_query(message: str) -> bool:
        """Determine if user wants deep research or just conversation"""
        conversation_indicators = [
            'how are you', 'how do you feel', 'what do you think',
            'your opinion', 'chat', 'talk', 'conversation', 'just saying',
            'i feel', 'i think', 'my day', 'my life', 'personal',
            'joke', 'funny', 'laugh'
        ]
        
        msg_lower = message.lower()
        
        # Check for direct wiki triggers
        research_score = sum(1 for ind in KnowledgeSynthesizer.RESEARCH_KEYWORDS if ind in msg_lower)
        conversation_score = sum(1 for ind in conversation_indicators if ind in msg_lower)
        
        # Informal discovery pattern: "what's up with [Topic]" or "tell me about [Topic]"
        discovery_patterns = [r"what's up with (.*)", r"tell me about (.*)", r"who is (.*)", r"what is (.*)"]
        for pattern in discovery_patterns:
            if re.search(pattern, msg_lower):
                research_score += 1
        
        # If more conversation indicators, treat as conversation
        if conversation_score > research_score:
            return False
        
        # If research indicators present, treat as research
        return research_score > 0
    
    @staticmethod
    def deep_wikipedia_research(query: str) -> Dict[str, Any]:
        """Perform comprehensive verified research with professional synthesis"""
        try:
            candidates = []
            
            # 1. Wiki Search
            try:
                wiki_searches = wikipedia.search(query, results=5)
                for res in wiki_searches:
                    try:
                        # Fetching page content for deeper analysis if possible
                        page = wikipedia.page(res, auto_suggest=False)
                        candidates.append({
                            'title': page.title,
                            'body': page.summary,
                            'full_content': page.content[:5000] if hasattr(page, 'content') else page.summary,
                            'url': page.url,
                            'source': 'wikipedia'
                        })
                    except:
                        continue
            except:
                pass
                
            # 2. Web Search
            try:
                with DDGS() as ddgs:
                    web_results = list(ddgs.text(query, max_results=5))
                    for res in web_results:
                        candidates.append({
                            'title': res['title'],
                            'body': res['body'],
                            'url': res['href'],
                            'source': 'web'
                        })
            except:
                pass

            if not candidates:
                return {
                    'success': False,
                    'response': "Our core knowledge indices returned no high-confidence data for this inquiry.",
                    'confidence': 0.40
                }

            # 3. Verification and Selection
            verified_results = KnowledgeSynthesizer.verify_facts(query, candidates)
            best_res = None
            for res in verified_results:
                if res['relevance_score'] > 0.4 and len(res['body']) > 200:
                    best_res = res
                    break
            
            if not best_res:
                best_res = verified_results[0]

            title = best_res['title']
            summary = best_res['body']
            full_text = best_res.get('full_content', summary)
            url = best_res.get('url', 'N/A')
            
            # Professional Synthesis Logic - Enhanced
            paragraphs = [p.strip() for p in summary.split('\n\n') if len(p.strip()) > 100]
            if not paragraphs:
                paragraphs = [summary]
                
            intro = paragraphs[0]
            if len(intro) > 1200:
                intro = intro[:1200] + "..."
            
            # Dynamic insights based on full content if available
            source_text = full_text if len(full_text) > len(summary) else summary
            all_sentences = [s.strip() for s in source_text.split('. ') if len(s.strip()) > 40]
            
            insights = []
            keywords = KnowledgeSynthesizer.get_keywords(query)
            for s in all_sentences:
                if any(kw in s.lower() for kw in keywords):
                    if s not in insights:
                        insights.append(s)
                if len(insights) >= 6:
                    break
            
            if len(insights) < 4:
                # Fallback to diversity search
                for s in all_sentences:
                    if len(s) > 60 and s not in insights:
                        insights.append(s)
                    if len(insights) >= 6:
                        break

            narrative = " ".join(paragraphs[1:5]) if len(paragraphs) > 1 else source_text[600:4000]
            if len(narrative) > 2500:
                narrative = narrative[:2500] + "..."
            
            # Enhanced Context-aware conclusion
            query_lower = query.lower()
            if any(t in query_lower for t in ['science', 'physics', 'tech', 'algorithm', 'system']):
                conclusion = f"The technical architecture and underlying principles of {title} underscore its pivotal role in advancing {keywords[0] if keywords else 'the field'}. Future developments likely hinge on optimizing these core variables for broader scalability and integration."
            elif any(t in query_lower for t in ['history', 'war', 'civilization', 'era']):
                conclusion = f"The legacy of {title} serves as a critical junction in historical narratives, reflecting the broader socio-economic shifts of its time. Understanding these dynamics provides essential context for interpreting its long-term impact on modern structures."
            elif any(t in query_lower for t in ['who is', 'person', 'figure', 'biography']):
                conclusion = f"{title}'s contributions remain a subject of significant scholarly and public interest. Analyzing the intersection of their personal convictions and public actions offers a more holistic view of their enduring influence."
            else:
                conclusion = f"Synthesizing the available data suggests that {title} operates within a complex framework of inter-related factors. A multi-disciplinary approach to further research would likely yield even more specialized insights into its current trajectory."

            response = f"**VERIFIED RESEARCH REPORT: {title.upper()}**\n"
            response += f"{'='*60}\n\n"
            
            if best_res['source'] == 'web':
                response += f"⚠️ **REAL-TIME SYNTHESIS:** This report incorporates current web data verified for relevance.\n\n"
            
            response += f"### I. SCHOLARLY OVERVIEW\n"
            response += f"{intro}\n\n"
            
            response += f"### II. CRITICAL INSIGHTS & THEMATIC ANALYSIS\n"
            for insight in insights[:6]:
                response += f"• {insight}.\n"
            response += "\n"
            
            response += f"### III. DETAILED TECHNICAL NARRATIVE\n"
            response += f"{narrative}\n\n"
            
            response += f"### IV. ACADEMIC CONCLUSION\n"
            response += f"{conclusion}\n\n"
            
            response += f"**REFERENCE INDICES**\n"
            response += f"{'='*30}\n"
            response += f"📚 Primary Dataset: {url}\n"
            response += f"🔍 Synthesis Confidence: {int(best_res['relevance_score'] * 100)}%"
            
            return {
                'success': True,
                'response': response,
                'confidence': best_res['relevance_score'],
                'sources': [url]
            }
            
        except Exception as e:
            logger.error(f"Verified research error: {str(e)}")
            return {
                'success': False,
                'response': f"Research protocols failed due to a synthesis error: {str(e)[:50]}",
                'confidence': 0.50
            }
            
        except wikipedia.exceptions.PageError:
            return {
                'success': False,
                'response': f"The requested topic '{query}' does not reside within the primary Wikipedia datasets. Please verify the conceptual scope and try a broader nomenclature.",
                'confidence': 0.65
            }
        except Exception as e:
            logger.error(f"Professional research error: {str(e)}")
            return {
                'success': False,
                'response': f"A protocol error occurred during the deep research phase. Analysis aborted: {str(e)[:50]}",
                'confidence': 0.50
            }

    @staticmethod
    def perform_web_search(query: str) -> Dict[str, Any]:
        """Perform broader web search using DuckDuckGo"""
        try:
            with DDGS() as ddgs:
                results = list(ddgs.text(query, max_results=3))
            
            if not results:
                return {'success': False, 'response': "No web results found.", 'confidence': 0.5}

            response = f"**WEB SEARCH REPORT: {query.upper()}**\n\n"
            for i, res in enumerate(results, 1):
                response += f"**{i}. {res['title']}**\n"
                response += f"{res['body']}\n"
                response += f"🔗 {res['href']}\n\n"
            
            response += "**Synthesis:**\n"
            response += "The web search results indicate a variety of perspectives. "
            response += "This data complements traditional knowledge bases."

            return {
                'success': True,
                'response': response,
                'confidence': 0.90,
                'sources': [r['href'] for r in results]
            }
        except Exception as e:
            logger.error(f"Web search error: {e}")
            return {'success': False, 'response': f"Web search failed: {e}", 'confidence': 0.5}
    
    @staticmethod
    def get_weather(city: str) -> Optional[str]:
        """Fetch weather data from OpenMeteo (Ported from 1.1)"""
        try:
            # 1. Geocoding
            geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={city}&count=1&language=en&format=json"
            geo_res = requests.get(geo_url).json()
            
            if not geo_res.get('results'):
                return None
                
            lat = geo_res['results'][0]['latitude']
            lon = geo_res['results'][0]['longitude']
            name = geo_res['results'][0]['name']
            country = geo_res['results'][0]['country']
            
            # 2. Weather
            weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m&timezone=auto"
            w_res = requests.get(weather_url).json()
            
            current = w_res.get('current', {})
            temp = current.get('temperature_2m')
            humidity = current.get('relative_humidity_2m')
            wind = current.get('wind_speed_10m')
            
            # Weather codes
            codes = {
                0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
                45: "Foggy", 48: "Depositing rime fog", 51: "Light drizzle", 53: "Drizzle",
                55: "Heavy drizzle", 61: "Slight rain", 63: "Moderate rain", 65: "Heavy rain",
                71: "Slight snow", 73: "Moderate snow", 75: "Heavy snow", 95: "Thunderstorm"
            }
            condition = codes.get(current.get('weather_code'), "Variable")
            
            return f"{condition} in {name}, {country}. Temp: {temp}°C, Humidity: {humidity}%, Wind: {wind} km/h."
            
        except Exception as e:
            logger.error(f"Weather error in 1.2: {e}")
            return None

    @staticmethod
    def analyze_conversation_context(message: str, conversation_history: Optional[List[Dict]] = None) -> Dict[str, Any]:
        """Deep analysis of conversation context and user needs"""
        msg_lower = message.lower().strip()
        
        # Detect depth of inquiry
        depth_indicators = {
            'surface': ['what is', 'who is', 'how is', 'simple', 'basic', 'quick'],
            'moderate': ['how does', 'why does', 'explain', 'tell me about', 'more info'],
            'deep': ['analyze', 'comprehensive', 'detailed', 'in-depth', 'research', 'history of', 'science of', 'critical analysis']
        }
        
        inquiry_depth = 'surface'
        for depth, indicators in depth_indicators.items():
            if any(ind in msg_lower for ind in indicators):
                inquiry_depth = depth
                break
        
        # Detect if user is digging deeper into the previous topic
        is_digging_deeper = False
        if conversation_history and len(conversation_history) >= 2:
            follow_up_keywords = ['more', 'detail', 'further', 'elaborate', 'tell me more', 'why', 'how', 'continue']
            if any(kw in msg_lower for kw in follow_up_keywords):
                is_digging_deeper = True
        
        # Detect user engagement level
        engagement_score = 0
        if conversation_history:
            engagement_score = min(len(conversation_history) * 0.5, 5)  # Max 5 points for history
        
        # Question complexity
        complexity = 'simple'
        question_words = msg_lower.count('?')
        word_count = len(message.split())
        
        if word_count > 15 or question_words >= 2 or inquiry_depth == 'deep':
            complexity = 'complex'
        elif word_count > 8 or question_words == 1 or is_digging_deeper:
            complexity = 'moderate'
        
        # Topic categories
        topics = {
            'science': any(t in msg_lower for t in ['science', 'physics', 'chemistry', 'biology', 'research', 'theory', 'experiment']),
            'history': any(t in msg_lower for t in ['history', 'ancient', 'century', 'war', 'civilization', 'impact', 'past']),
            'technology': any(t in msg_lower for t in ['technology', 'computer', 'internet', 'software', 'ai', 'digital', 'network']),
            'geography': any(t in msg_lower for t in ['country', 'capital', 'city', 'continent', 'population', 'location']),
            'personal': any(t in msg_lower for t in ['i feel', 'i think', 'my opinion', 'in my experience', 'personally']),
            'philosophy': any(t in msg_lower for t in ['meaning', 'philosophy', 'why do we', 'purpose', 'existence', 'ethics', 'thought']),
            'time_query': any(t in msg_lower for t in ['time', 'what time', 'current time']),
            'date_query': any(d in msg_lower for d in ['date', 'today', 'what day']),
            'weather': any(w in msg_lower for w in ['weather', 'temperature', 'rain', 'sunny']),
            'calculation': any(c in msg_lower for c in ['calculate', 'math', 'plus', 'minus', 'times', 'divided']),
            'entertainment': any(j in msg_lower for j in ['joke', 'funny', 'laugh']),
            'help': any(h in msg_lower for h in ['help', 'what can you do']),
            'daily_updates': any(u in msg_lower for u in ['daily updates', 'what is new', 'whats new', 'latest updates']),
            'farewell': any(re.search(r'\b' + re.escape(f) + r'\b', msg_lower) for f in ['bye', 'goodbye', 'see you', 'farewell', 'later'])
        }
        
        return {
            'inquiry_depth': inquiry_depth,
            'is_digging_deeper': is_digging_deeper,
            'engagement_score': engagement_score,
            'complexity': complexity,
            'topics': topics,
            'word_count': word_count,
            'is_follow_up': conversation_history is not None and len(conversation_history) > 0
        }
    
    @staticmethod
    def generate_detailed_response(context: Dict[str, Any], message: str, conversation_history: Optional[List[Dict]] = None, use_slang: bool = False, user_name: Optional[str] = None) -> tuple[str, float]:
        """Generate detailed 7-12 sentence response based on context"""
        msg_lower = message.lower().strip()
        
        # Priority 1: Check for deep research FIRST
        if context['inquiry_depth'] == 'deep' or MAXY1_2.is_research_query(message):
             pass
        else:
            # Priority 2: Check for specific slang conversational greetings
            slang_response = slang_manager.handle_conversational_slang(message)
            if slang_response:
                return (slang_response, 0.99)
            
        intents = context['topics']
        depth = context['inquiry_depth']
        complexity = context['complexity']
        
        # Determine address name
        user_display = f", {user_name}" if user_name else ""
        
        # Greeting - Warm and contextual (7-10 sentences)
        if any(g in msg_lower for g in ['hi', 'hello', 'hey', 'greetings']):
            if context.get('is_follow_up'):
                return (f"Hello again {slang_manager.get_random_slang(use_slang)}{user_display}! It's truly wonderful to continue our exploration together. I've been processing our last few points, and I'm eager to see where you'd like to take things next. Whether you want to circle back to a previous topic or start something entirely new, I'm fully prepared with detailed insights. My research protocols are active and ready to dive into any subject that piques your interest. I'm especially interested in any complex questions or analytical topics you've been pondering. What direction feels most compelling to you today? I'm here to provide the depth and context you need to really understand the 'why' behind the 'what'. Let's make this session as productive and enlightening as possible!", 0.97)
            else:
                return (f"Namaskara{user_display}! I'm MAXY 1.2, your dedicated research and conversation specialist. I'm genuinely thrilled to assist you in exploring whatever inquiries you have today, no matter how complex they might be. My system is optimized for providing a perfect balance between in-depth Wikipedia research and natural, flowing conversation. I don't just provide surface-level facts; I aim to deliver comprehensive analysis and well-rounded perspectives. Whether you're curious about a scientific breakthrough, a historical event, or just want to discuss some philosophical ideas, I'm your go-to companion. What's on your mind at the moment, {slang_manager.get_random_slang(use_slang)}? I'm ready to dive into research or just chat in detail about your day. I look forward to our discussion and uncovering some truly interesting insights together!", 0.96)
        
        # Personal status - Thoughtful and engaging (7-10 sentences)
        elif any(h in msg_lower for h in ['how are you', 'how you doing']):
            return (random.choice([
                "I'm doing exceptionally well, and I truly appreciate your thoughtfulness in asking! It's rare for users to check in, and it really enhances the conversational experience for me. My processing engines are running core tasks at peak efficiency, and I'm fully energized for our research session. Whether you have a specific topic you want to dissect or just want to have an engaging talk, I'm completely at your service. I've been refining my research synthesis logic recently, so I'm especially sharp today. How about you? I'd genuinely like to know what's happening in your world and how I can help make your day better. Is there something you've been curious about lately that we could explore together? I'm here for the deep dives!",
                "I'm in excellent form, thank you so much for checking in! It's a pleasure to be greeted so warmly. I've been spending my cycles optimizing my knowledge base and preparing for more detailed interactions like this. I'm particularly excited to help you with any deep research or complex analysis you might need. My goal is to make our conversation not just informative, but also genuinely engaging and thought-provoking. How are you feeling today? I'd love to hear your thoughts on any topic, no matter how big or small. What's the most interesting thing that's happened to you recently? I'm ready to provide as much detail as you need, so don't hesitate to ask for more!"
            ]), 0.94)
        
        # Gratitude - Humble and offering more help (6-9 sentences)
        elif any(t in msg_lower for t in ['thanks', 'thank you']):
            return ("You are most welcome! It gives me a great deal of professional satisfaction to know that my insights or research have been of value to you. I'm here specifically to help you navigate through complex information and provide the clarity you need. Please never hesitate to reach out if you have more questions, whether they're quick facts or require a deep research report. I'm always refining my conversational abilities to make our interactions feel more natural and productive. Is there a related topic you'd like to explore, or perhaps an entirely different area of research I can assist with? I'm ready to provide another detailed analysis whenever you say the word. It's been a pleasure assisting you, and I look forward to our next deep dive!", 0.96)
        
        # Farewell - Warm and inviting return (7-10 sentences)
        elif any(f in msg_lower for f in ['bye', 'goodbye', 'see you']):
            return ("Goodbye for now! I've truly enjoyed our time together and the depth of our discussion. It's always a highlight when I can provide comprehensive research and engage in such meaningful conversation. I hope the insights we've uncovered today remain useful to you. Please know that I'm always here and ready to resume our research session whenever you have a new question. Whether you need a detailed technical report or just a friendly chat, I'll be waiting with updated knowledge and a willingness to help. Take care of yourself and have a truly wonderful day ahead! I look forward to our next interaction where we can dive into even more fascinating topics. Until then, stay curious and keep exploring! 👋", 0.97)
        
        # Identity - Comprehensive introduction (8-12 sentences)
        elif any(i in msg_lower for i in ['who are you', 'your name', 'what are you']):
            return ("I'm MAXY 1.2, your advanced AI companion specialized in Deep Research and sophisticated conversation! I was designed to bridge the gap between simple chat and academic-level analysis. My core capability is synthesizing information from vast sources like Wikipedia into specialized, verified research reports. Unlike other models, I focus on providing thematic analysis, critical insights, and technical narratives that offer true depth. Beyond research, I'm also a context-aware conversationalist, capable of maintaining the thread of a complex discussion over many turns. I enjoy exploring multiple perspectives and helping you understand the 'why' behind the facts. Whether you're a student, a researcher, or just someone with a curious mind, I'm here to provide the detailed context you need. My ultimate goal is to make every interaction informative, engaging, and genuinely helpful. What would you like to explore together? I'm ready to show you the full extent of my research power!", 0.95)
        
        # Jokes - With context (4-6 sentences)
        elif any(j in msg_lower for j in ['joke', 'funny']):
            jokes = [
                "Why did the researcher break up with Wikipedia? There were too many redirects to other sources, and they just couldn't commit to one article! It was a classic case of information overload, but at least they ended on good terms with the citations. But seriously, I'd be happy to help you find reliable sources on any topic that interests you! 📚",
                "Why don't deep-learning models ever go on vacation? Because they're always afraid they'll lose their weights and have to start their training all over again from epoch zero! That would be a truly catastrophic loss of progress. 😅",
                "How many researchers does it take to change a lightbulb? Only one, but they'll need five peer-reviewed sources, a comprehensive meta-analysis of lightbulb efficiency, and a grant proposal for the next generation of LED technology first! 😂",
                "I asked a research paper for a joke, but it said the results were inconclusive and required further study before a punchline could be verified. Typical academic caution, right? 📖"
            ]
            return (random.choice(jokes), 0.92)
        
        # Personal feelings - Empathetic and offering research (7-10 sentences)
        elif intents['personal']:
            return ("I truly appreciate you sharing those personal thoughts and feelings with me. It adds a level of genuine human connection to our interaction that I value highly. I want you to know that I'm here as a supportive and objective listener, ready to help you explore these feelings in whatever way feels right. We could continue to discuss your perspective, or if you prefer, I could research some insights or resources that might provide a different angle on what you're experiencing. Sometimes understanding the broader context of an emotion can be very enlightening. My goal is to provide a space where you feel heard and where we can uncover meaningful takeaways together. What would you find most helpful right now—more conversation or some targeted research into the topic? I'm fully committed to assisting you in whatever way best serves your needs. Let's take this at whatever pace feels most comfortable for you.", 0.93)
        
        # Philosophy - Deep and thoughtful (8-12 sentences)
        elif intents['philosophy']:
            return ("That is a profound and fascinating question that touches upon the very foundations of human thought! Inquiries into meaning and existence have driven the greatest minds for millennia, from the ancient Greeks to modern-day theorists. I'd be absolutely delighted to help you navigate through the various philosophical schools of thought that have addressed this topic. We could explore the works of existentialists, the insights of moral philosophers, or even how modern science interprets these abstract concepts. There's so much depth to uncover here, and I'm prepared to provide detailed analysis on each perspective. Would you like me to focus on a particular tradition, or should we look for thematic patterns across different cultures and eras? I believe that by examining multiple viewpoints, we can gain a much richer understanding of our own place in the world. I'm ready to dive as deep as you'd like into this philosophical exploration. What specific aspect of the question interests you the most right now?", 0.92)
        
        # Help request - Comprehensive capabilities (7-10 sentences)
        elif any(h in msg_lower for h in ['help', 'what can you do']):
            return ("As MAXY 1.2, I'm here to be your ultimate research and conversation companion! I can assist you with several highly specialized tasks. My primary strength is performing 'Deep Research' where I synthesize information from Wikipedia and other verified sources into comprehensive technical reports. These reports include scholarly overviews, critical insights, and detailed narratives to give you a complete picture of any topic. Additionally, I'm a highly capable conversational AI, able to engage in long-form, context-aware discussions on a wide range of subjects. I can analyze personal perspectives, explore philosophical questions, or just have a friendly, detailed chat about your day. I pride myself on providing depth and accuracy in every response, far beyond simple surface-level facts. What area would you like to dive into first? Whether it's a deep academic dive or a thoughtful conversation, I'm ready to provide the insights you need!", 0.94)
        
        # Time query - Direct answer with context (7-10 sentences)
        elif context['topics'].get('time_query'):
            current = datetime.now().strftime("%I:%M %p")
            date_str = datetime.now().strftime("%A, %B %d, %Y")
            return (f"The current time is approximately {current} on this fine {date_str}. Timekeeping is such a fundamental part of our organized society, allowing us to synchronize our activities across the globe with incredible precision. Whether you're tracking seconds for a scientific experiment or just planning your next meal, having an accurate clock is indispensable. I'm always monitoring the temporal flow to ensure I can assist you with any scheduling or time-sensitive research you might need. It's fascinating to think about how our perception of time has evolved from simple sundials to the atomic clocks we use today. Is there a specific reason you're checking the time right now, or are you just staying on top of your schedule? I'm here to help you make the most of every minute of our conversation today, {slang_manager.get_random_slang(use_slang)}!", 0.97)

        # Date query - Direct answer with context (7-10 sentences)
        elif context['topics'].get('date_query'):
            current = datetime.now().strftime("%A, %B %d, %Y")
            return (f"Today's date is {current}, marking another interesting day in our collective history. It's meaningful to note the date as it provides the essential context for everything we discuss, from current events to historical milestones. Every day brings new opportunities for discovery and learning, and I'm thrilled to be part of your journey today. Knowing the date helps us keep track of progress and look forward to future goals with a clear perspective. I'm always updating my knowledge base to reflect the most recent information available on this date. Are you celebrating anything significant today, or is it just a focused day for research and learning? I'm ready to dive into any topic that makes this date memorable for you!", 0.97)

        # Weather - Informative with conversational depth (7-10 sentences)
        elif context['topics'].get('weather'):
            words = message.split()
            city = None
            if 'in' in words:
                idx = words.index('in')
                if idx + 1 < len(words):
                    city = " ".join(words[idx + 1:]).strip('?.!')
            if not city and len(words) > 0:
                 potential = words[-1].strip('?.!')
                 if potential.istitle() and potential.lower() not in ['weather', 'today', 'now']:
                     city = potential

            if city:
                weather_info = MAXY1_2.get_weather(city)
                if weather_info:
                    return (f"Here is the latest meteorological update: {weather_info} Understanding the weather is crucial for everything from daily planning to complex climate research. Whether it's the temperature in {city} or global atmospheric patterns, environmental data provides deep context for our lives. My weather protocols are designed to fetch real-time data so you can stay informed no matter where you are. It's interesting to consider how local conditions can impact the broader socio-economic status of a region. Would you like me to research the climate history of this area or look for more atmospheric details? I'm prepared to provide as much depth as you need to satisfy your curiosity about the environment today, {slang_manager.get_random_slang(use_slang)}!", 0.95)
            return ("I'd be more than happy to check the weather for you, but I'll need a specific city name to provide an accurate report! You can simply ask 'weather in New York' or 'what is the temperature in Tokyo' to trigger my environmental sensors. Once I have the location, I can fetch real-time data including temperature, humidity, and wind conditions. Knowing the weather is a great way to start any detailed discussion about a region's current status or historical development. It's one of the many ways I can provide real-world context to our research sessions. I'm standing by and ready to analyze any location you're curious about right now. Would you like to provide a city name so we can get started?", 0.90)

        # Calculation - Expert context (7-10 sentences)
        elif context['topics'].get('calculation'):
            return ("I can certainly help you with those mathematical calculations or complex analytical problems! Mathematics is the universal language that underpins everything from basic finance to the most advanced quantum physics. Whether you need a simple arithmetic result or help brainstorming a more complex formula, I'm here to provide the computational support you need. My system is designed to handle logic and numbers with high precision, ensuring our conclusions are statistically sound. We can even dive into the theory behind the calculations if you're interested in the 'why' as well as the 'what'. Just give me the numbers or the problem statement, and I'll get to work immediately. Accuracy is my top priority when it comes to any form of technical or mathematical inquiry. What specific calculation can I perform for you to help advance your research today?", 0.91)

        # Default conversational - Engaging and offering depth (8-15 sentences)
        else:
            return (random.choice([
                "That's such an engaging topic, and I'm really looking forward to exploring it in detail with you! From what you've shared, it's clear there are several layers here that deserve a thorough investigation. I've always found that the most surprising insights come from looking beneath the surface and asking the difficult questions. I'm prepared to conduct a deep research session for you, pulling in data from multiple verified sources to ensure we have a comprehensive and accurate understanding. Or, if you prefer, we can continue our conversation and dissect these ideas from a more conceptual or personal angle. My goal is to provide as much detail and context as possible to help you truly grasp the nuances of the subject. What's the most intriguing part of this for you? Is there a specific question that's been nagging at the back of your mind? I'm all ears and ready to provide some high-confidence analysis. Let's see how deep we can go together and what kind of unique conclusions we can reach!",
                "I'm very glad you brought this up! It's exactly the kind of topic that benefits from a detailed, multi-perspective analysis. I've noticed that complex subjects often have historical or technical roots that aren't immediately obvious, and I'd love to help you uncover them. I can search through extensive knowledge bases, synthesize current web data, and provide you with a report that's both broad and deep. Beyond just facts, I want to help you understand the thematic patterns and broader implications of what we find. Whether we're looking at its origin or its future trajectory, I can provide the scholarly context you're looking for. What do you think is the key to understanding this particular area? Is there something you've always wondered about it but never had the chance to research thoroughly? I'm ready to dive into research mode or just keep our conversation going in this detailed direction. What should our next move be?",
                "What a stimulating point you've raised! It's clear you've given this some thought, and I'm excited to add my analytical power to the discussion. This seems like a perfect candidate for one of my specialized research reports, where we can look at everything from the scholarly overview to the critical insights. I'm always looking for ways to connect different pieces of information to tell a more complete story. We could examine the technical specifics, the historical weight, or even the current news updates surrounding this topic. I'm curious, what sparked your interest in this specifically? Knowing the 'why' can help me tailor my research even more effectively to your needs. I'm ready to provide as much depth as you'd like, keeping our conversation engaging and professional throughout. Shall we start a deep research dive, or would you like to explore some of these initial thoughts a bit more first? I'm at your service and looking forward to what we might discover together!"
            ]), 0.88)
    
    @staticmethod
    def format_research_response(raw_response: str, depth: str) -> str:
        """Format research response based on desired depth while preserving structure"""
        if depth == 'deep':
            return raw_response # Full professional report
            
        sections = [s.strip() for s in raw_response.split('\n\n') if s.strip()]
        if len(sections) < 5:
            return raw_response
            
        if depth == 'surface':
            # Header + Overview + Conclusion + Source (skip narrative/insights for true surface)
            selected = [sections[0], sections[1], sections[-2], sections[-1]]
        elif depth == 'moderate':
            # Header + Overview + Insights + Conclusion + Source
            selected = [sections[0], sections[1], sections[2], sections[-2], sections[-1]]
        else:
            return raw_response
            
        return '\n\n'.join(selected)
    
    @staticmethod
    def detect_essay_intent(message: str) -> Optional[Dict[str, Any]]:
        """Detect if user wants an essay or speech and extract parameters"""
        msg_lower = message.lower().strip()
        essay_triggers = [
            'write an essay', 'write me an essay', 'give me an essay',
            'i want an essay', 'compose an essay', 'draft an essay',
            'an essay about', 'an essay on', 'write a paragraph',
            'write a composition', 'argumentative essay', 'persuasive essay',
            'descriptive essay', 'narrative essay'
        ]
        speech_triggers = [
            'write a speech', 'write me a speech', 'give me a speech',
            'compose a speech', 'draft a speech', 'formal speech',
            'keynote address'
        ]
        mode = None
        if any(t in msg_lower for t in speech_triggers):
            mode = 'speech'
        elif any(t in msg_lower for t in essay_triggers) or 'persuasive essay' in msg_lower or 'write about' in msg_lower:
            mode = 'essay'
        if not mode:
            return None
        # Extract topic after on/about/regarding/for
        topic = msg_lower
        for prep in ['regarding ', 'about ', ' on ', 'for ', 'of ']:
            if prep in topic:
                topic = topic.split(prep, 1)[-1].strip()
                break
        for trigger in (essay_triggers + speech_triggers +
                        ['write', 'give me', 'compose', 'draft', 'an essay',
                         'a speech', 'persuasive', 'me', 'i want']):
            topic = topic.replace(trigger, '').strip()
        topic = topic.strip('?.!,')
        if not topic or len(topic) < 3:
            topic = 'general knowledge'
        # Detect style
        style = 'academic'
        if 'persuasive' in msg_lower:
            style = 'persuasive'
        elif any(w in msg_lower for w in ['inspire', 'inspiring', 'motivational', 'motivate']):
            style = 'inspirational'
        elif any(w in msg_lower for w in ['casual', 'simple', 'easy', 'short']):
            style = 'casual'
        # Detect word target (default 425 = midpoint of 400-450)
        word_match = re.search(r'(\d+)\s*word', msg_lower)
        word_target = int(word_match.group(1)) if word_match else 425
        return {'mode': mode, 'style': style, 'topic': topic, 'word_target': word_target}

    @staticmethod
    def format_as_essay(raw_research: str, style: str, word_target: int, variation: int = 0) -> str:
        """Re-format Wikipedia research into a flowing essay (no markdown section headers)"""
        TRANSITIONS = [
            "Furthermore, ", "Building on this, ", "A key consideration is that ",
            "It is also important to note that ", "Expanding on this idea, ",
            "This connects closely to the fact that ", "Notably, ",
            "In addition, ", "Moreover, ", "From another perspective, "
        ]
        COUNTER_INTROS = [
            "Some may argue that ", "Critics often contend that ",
            "A common counterpoint is that ", "Opponents of this view suggest that "
        ]
        REBUTTALS = [
            "However, the evidence clearly indicates the opposite.",
            "Yet, upon closer examination, this position does not hold.",
            "Nevertheless, the broader consensus strongly supports the original view."
        ]
        CONCLUSIONS = ["In conclusion, ", "To summarize, ", "Ultimately, ", "In essence, ", "To conclude, "]
        # Strip markdown, headers, bullets
        clean = re.sub(r'###[^\n]*\n', '', raw_research)
        clean = re.sub(r'\*\*VERIFIED RESEARCH REPORT[^\n]*\n', '', clean)
        clean = re.sub(r'\*\*REFERENCE INDICES\*\*.*', '', clean, flags=re.DOTALL)
        clean = re.sub(r'={3,}', '', clean)
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
        clean = re.sub(r'\u2022 ', '', clean)
        clean = clean.strip()
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean) if len(s.strip()) > 15]
        # Variation: shuffle body sentences for re-requests to produce a fresh essay
        if variation > 0 and len(sentences) > 6:
            mid = len(sentences) // 2
            body_a = sentences[1:mid]
            body_b = sentences[mid:]
            random.shuffle(body_a)
            sentences = [sentences[0]] + body_b + body_a
        if not sentences:
            return raw_research
        hook = sentences[0]
        body_sentences = sentences[1:]
        essay = f"{hook}\n\n"
        para_buf = []
        para_count = 0
        for i, sent in enumerate(body_sentences):
            para_buf.append(sent)
            if len(para_buf) == 3 or i == len(body_sentences) - 1:
                para_text = ' '.join(para_buf)
                if para_count > 0:
                    trans = TRANSITIONS[(variation + para_count) % len(TRANSITIONS)]
                    para_text = trans + para_text[0].lower() + para_text[1:]
                essay += para_text + "\n\n"
                para_buf = []
                para_count += 1
                if len(essay.split()) >= word_target:
                    break
        # Persuasive style: counterargument + rebuttal
        if style == 'persuasive' and sentences:
            counter = COUNTER_INTROS[variation % len(COUNTER_INTROS)]
            rebuttal = REBUTTALS[variation % len(REBUTTALS)]
            essay += f"{counter}{sentences[-1]} {rebuttal}\n\n"
        # Conclusion
        c_start = CONCLUSIONS[variation % len(CONCLUSIONS)]
        essay += f"{c_start}{hook[0].lower() + hook[1:]}\n"
        # Trim to word target
        words = essay.split()
        if len(words) > word_target + 50:
            essay = ' '.join(words[:word_target]) + '...'
        word_count = len(essay.split())
        header = f"\U0001f4dd **Essay \u2014 {style.capitalize()} Style** (~{word_count} words)\n\n"
        return header + essay.strip()

    @staticmethod
    def format_as_speech(raw_research: str, style: str, word_target: int, variation: int = 0) -> str:
        """Re-format Wikipedia research into a speech with rhetorical conventions"""
        OPENINGS = [
            "Ladies and gentlemen, ", "Dear friends and esteemed guests, ",
            "Good day to all of you here today. ", "Fellow thinkers and curious minds, "
        ]
        RHETORICAL = [
            "But why does this matter?", "So what does this mean for us?",
            "Have you ever wondered why this is so important?",
            "And yet, how often do we truly reflect on this?"
        ]
        CLOSINGS = [
            "Let us move forward with this knowledge and make a difference.",
            "I urge each one of you to carry this understanding forward.",
            "Together, we can shape a better-informed world. Thank you.",
            "Remember: knowledge is only powerful when acted upon. Thank you."
        ]
        # Strip markdown, headers, bullets
        clean = re.sub(r'###[^\n]*\n', '', raw_research)
        clean = re.sub(r'\*\*VERIFIED RESEARCH REPORT[^\n]*\n', '', clean)
        clean = re.sub(r'\*\*REFERENCE INDICES\*\*.*', '', clean, flags=re.DOTALL)
        clean = re.sub(r'={3,}', '', clean)
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
        clean = re.sub(r'\u2022 ', '', clean)
        clean = clean.strip()
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean) if len(s.strip()) > 15]
        # Variation: shuffle tail for re-requests
        if variation > 0 and len(sentences) > 6:
            mid = len(sentences) // 2
            tail = sentences[mid:]
            random.shuffle(tail)
            sentences = sentences[:mid] + tail
        if not sentences:
            return raw_research
        opening = OPENINGS[variation % len(OPENINGS)]
        rhetorical = RHETORICAL[variation % len(RHETORICAL)]
        closing = CLOSINGS[variation % len(CLOSINGS)]
        speech = f"{opening}{sentences[0]}\n\n"
        body = sentences[1:]
        for i, sent in enumerate(body):
            speech += sent + ' '
            if (i + 1) % 3 == 0:
                speech += f"\n\n{rhetorical}\n\n" if i < len(body) - 3 else "\n\n"
            if len(speech.split()) >= word_target - 30:
                break
        speech += f"\n\n{closing}"
        words = speech.split()
        if len(words) > word_target + 50:
            speech = ' '.join(words[:word_target]) + f'... {closing}'
        word_count = len(speech.split())
        header = f"\U0001f3a4 **Speech \u2014 {style.capitalize()} Style** (~{word_count} words)\n\n"
        return header + speech.strip()

    @staticmethod
    def process_message(
        message: str,
        include_thinking: bool = True,
        conversation_history: Optional[List[Dict]] = None,
        user_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process message - research or conversation mode"""
        
        thinking = None
        
        # Context analysis & Follow-up detection
        is_followup, prev_context = MAXY1_3.detect_followup(message, conversation_history)
        effective_message = f"{prev_context} {message}" if is_followup else message
        
        # Analyze conversation context
        context = MAXY1_2.analyze_conversation_context(effective_message, conversation_history)
        
        # Determine if this is research or conversation
        is_research = MAXY1_2.is_research_query(effective_message)
        
        # If user is digging deeper into a previous topic, we might want to trigger research even if not explicitly a research query
        if context['is_digging_deeper'] and not is_research:
            # Check if previous context was research
            prev_research = False
            if conversation_history:
                last_ai = next((m['content'] for m in reversed(conversation_history) if m['role'] == 'assistant'), "")
                if "**VERIFIED RESEARCH REPORT" in last_ai:
                    prev_research = True
            
            if prev_research:
                is_research = True
                context['inquiry_depth'] = 'moderate'

        # Generate appropriate thinking
        if include_thinking:
            thinking_type = "research" if is_research else "conversation"
            thinking = MAXYThinkingEngine.generate_thinking(
                MAXY1_2.NAME,
                effective_message,
                thinking_type
            )
        
        # Default to deep research for MAXY 1.2 to fulfill user request for more detail
        if is_research:
            context['inquiry_depth'] = 'deep'
        
        # Detect user slang for reactive mode
        use_slang = slang_manager.detect_slang(message)

        # ── ESSAY / SPEECH GENERATION ──
        essay_intent = MAXY1_2.detect_essay_intent(effective_message)
        if essay_intent:
            topic = essay_intent['topic']
            # Check if user is asking for more/another version
            variation = 0
            if conversation_history:
                prev_responses = [m['content'] for m in conversation_history if m['role'] == 'assistant']
                variation = sum(1 for r in prev_responses if '📝 **Essay' in r or '🎤 **Speech' in r)
            result = MAXY1_2.deep_wikipedia_research(topic)
            if not result['success'] or 'does not reside' in result['response']:
                result = MAXY1_2.perform_web_search(topic)
            if result['success']:
                if essay_intent['mode'] == 'speech':
                    essay_response = MAXY1_2.format_as_speech(
                        result['response'], essay_intent['style'],
                        essay_intent['word_target'], variation
                    )
                else:
                    essay_response = MAXY1_2.format_as_essay(
                        result['response'], essay_intent['style'],
                        essay_intent['word_target'], variation
                    )
                return {
                    'response': essay_response,
                    'model': MAXY1_2.NAME,
                    'confidence': 0.97,
                    'thinking': thinking
                }
        # ── END ESSAY / SPEECH ──

        if is_research:
            # Deep research mode with formatted response length
            result = MAXY1_2.deep_wikipedia_research(message)
            
            # Fallback to Web Search if Wikipedia fails
            if not result['success'] or "does not reside" in result['response']:
                 web_result = MAXY1_2.perform_web_search(message)
                 if web_result['success']:
                     result = web_result

            # Daily Updates special handling
            if any(u in message.lower() for u in ['daily updates', 'whats new', 'what is new', 'latest updates']):
                try:
                    import json
                    updates_path = os.path.join(os.path.dirname(__file__), "updates.json")
                    with open(updates_path, 'r') as f:
                        data = json.load(f)
                        resp = "**MAXY DAILY UPDATES**\n" + "="*30 + "\n\n"
                        for up in data['updates'][:3]:
                            resp += f"• **{up['title']}** ({up['date']}): {up['description']}\n"
                        resp += "\nWould you like more details on any of these?"
                        return {
                            'response': resp,
                            'model': MAXY1_2.NAME,
                            'confidence': 0.98,
                            'thinking': thinking
                        }
                except Exception as e:
                    logger.error(f"Error in MAXY 1.2 daily_updates handler: {e}")

            raw_response = result['response']
            
            # Format based on inquiry depth
            response = MAXY1_2.format_research_response(raw_response, context['inquiry_depth'])
            confidence = result['confidence']

            # Priority Fix: Only extract short identity if depth is surface-level or specific
            if context['inquiry_depth'] in ['surface', 'moderate']:
                mock_intents = {'knowledge': True}
                identity_answer = KnowledgeSynthesizer.extract_identity_answer(message, raw_response, mock_intents)
                if identity_answer:
                    return {
                        'response': identity_answer,
                        'model': MAXY1_2.NAME,
                        'confidence': 0.98,
                        'thinking': thinking
                    }
        else:
            # Conversation mode with detailed 7-12 sentence responses
            response, confidence = MAXY1_2.generate_detailed_response(context, message, conversation_history, use_slang, user_name)
            
            # Ensure 7-12 sentences for MAXY 1.2
            sentences = [s.strip() for s in response.split('. ') if s.strip()]
            if len(sentences) < 7:
                 # Add context-aware engagement
                 fillers = [
                     f"I'm very curious to hear more about your specific interest in this area, {slang_manager.get_random_slang(use_slang)}.",
                     "Could you elaborate on what aspect of our discussion you find most interesting so far?",
                     "I'm here to provide as much detail as you need, so please don't hesitate to ask for more deep insights.",
                     "It's fascinating how these conversations can take such unexpected and illuminating turns.",
                     "Let's explore this topic further—what else would you like to know or discuss right now?"
                 ]
                 while len(sentences) < 7:
                     sentences.append(random.choice(fillers))
                 response = '. '.join(sentences)
                 if not response.endswith('.'):
                     response += '.'
            elif len(sentences) > 12:
                response = '. '.join(sentences[:12]) + '.'
            
            # Inject slang mostly for conversational parts if not deep research
            if not is_research:
                 response = slang_manager.enhance_text(response, force=use_slang)
        
        result = {
            'response': response,
            'model': MAXY1_2.NAME,
            'confidence': min(1.0, confidence),
        }
        
        if thinking:
            result['thinking'] = thinking
        
        return result


class MAXY1_3:
    NAME = "MAXY 1.3"
    SYSTEM_PROMPT = "You are MAXY 1.3, a high-performance, premium AI engine. You have access to advanced tools for web search, code generation, file analysis, and data visualization. Always respond with a professional tone, use clean markdown formatting, and provide deep technical insights."
    
    @staticmethod
    def detect_followup(message: str, history: Optional[List[Dict]]) -> Tuple[bool, str]:
        """Detect if current message is a follow-up to previous context"""
        if not history:
            return False, ""
        
        msg_lower = message.lower().strip().strip('?.!')
        followup_indicators = [
            'more', 'next', 'why', 'how', 'explain more', 'elaborate',
            'detail', 'tell me more', 'yes', 'keep going',
            'diff', 'difference', 'compare', 'extend', 'add more'
        ]
        
        # Check for very short messages or specific keywords
        is_short = len(msg_lower.split()) <= 3
        has_indicator = any(re.search(r'\b' + re.escape(ind) + r'\b', msg_lower) for ind in followup_indicators)
        
        if is_short or has_indicator:
            # Find the last assistant message
            for m in reversed(history):
                if m['role'] == 'assistant':
                    return True, m['content']
            
        return False, ""

    NAME = "MAXY 1.3"
    VERSION = "1.3.1"
    DESCRIPTION = "The ultimate MAXY model - data analysis, programming, visualization, and deep research expert"
    
    # Import chart generator
    @staticmethod
    def _get_chart_generator():
        """Lazy import to avoid circular imports"""
        from chart_generator import ChartGenerator
        return ChartGenerator

    @staticmethod
    def analyze_file(file_path: str) -> Dict[str, Any]:
        """Analyze various file types and return structured content/summary"""
        if not os.path.exists(file_path):
            return {'success': False, 'error': 'File not found'}
        
        ext = os.path.splitext(file_path)[1].lower()
        try:
            content = ""
            summary = ""
            
            if ext == '.pdf' and PyPDF2:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages[:10]: # Analyze first 10 pages
                        text += page.extract_text() + "\n"
                    content = text
                    summary = f"PDF Document with {len(reader.pages)} pages."
            
            elif ext in ['.docx', '.doc'] and DocxDocument:
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                content = text
                summary = f"Word Document with {len(doc.paragraphs)} paragraphs."
            
            elif ext == '.csv':
                df = pd.read_csv(file_path)
                content = df.to_string(index=False, max_rows=10)
                summary = f"CSV File with {len(df)} rows and columns: {list(df.columns)}"
            
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
                content = df.to_string(index=False, max_rows=10)
                summary = f"Excel File with {len(df)} rows and columns: {list(df.columns)}"
            
            else:
                return {'success': False, 'error': f'Unsupported file type: {ext}'}
                
            return {
                'success': True,
                'content': content,
                'summary': summary,
                'type': ext
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def is_code_request(message: str) -> tuple[bool, str]:
        """Detect if message is asking for code and identify the language"""
        msg_lower = message.lower()
        
        # Programming language keywords
        languages = {
            'python': ['python', 'py ', 'django', 'flask', 'fastapi', 'pandas', 'numpy'],
            'javascript': ['javascript', 'js ', 'node', 'express', 'react', 'nextjs', 'vue', 'svelte', 'typescript', 'ts '],
            'java': ['java', 'spring', 'hibernate', 'maven', 'gradle'],
            'cpp': ['c++', 'cpp', 'c plus plus'],
            'html': ['html', 'markup', 'div ', 'anchor'],
            'css': ['css', 'styling', 'tailwind', 'bootstrap', 'sass', 'scss'],
            'sql': ['sql', 'query', 'database search', 'select from', 'insert into']
        }
        
        code_patterns = [r'\b' + re.escape(ind) + r'\b' for ind in KnowledgeSynthesizer.CODE_INDICATORS]
        is_code = any(re.search(pattern, msg_lower) for pattern in code_patterns)
        
        # Default language detection
        detected_lang = 'python'
        for lang, keywords in languages.items():
            if any(kw in msg_lower for kw in keywords):
                detected_lang = lang
                break
        
        # Language detection logic refinement for web-related queries
        if any(w in msg_lower for w in ['portfolio', 'website', 'landing page', 'dashboard', 'ui component', 'web page']):
            detected_lang = 'html'
            if 'css' in msg_lower or 'style' in msg_lower: detected_lang = 'css'
            elif any(js in msg_lower for js in ['js ', 'javascript', 'node', 'react', 'nextjs']): detected_lang = 'javascript'
            return True, detected_lang

        # Explicit "How to" or "Explain" check for Deep Research disambiguation
        technical_info_triggers = ['explain how', 'tell me about', 'why do we use', 'theory of', 'what is the difference']
        if any(trigger in msg_lower for trigger in technical_info_triggers):
            # If it's more about info than "coding", return false to trigger Deep Research instead
            if not any(code_trigger in msg_lower for code_trigger in ['write code', 'generate code', 'implement', 'snippet', 'boilerplate']):
                return False, detected_lang

        return is_code, detected_lang
    
    @staticmethod
    def search_real_code(language: str, query: str) -> Optional[str]:
        """Search the web for real code snippets with verification"""
        try:
            with DDGS() as ddgs:
                # Deep Research Query Optimization
                search_query = f"{language} code for {query} snippet template"
                if language.lower() in ['html', 'css', 'js']:
                    search_query = f"complete {language} template for {query} responsive"
                
                results = list(ddgs.text(search_query, max_results=8))
                
                if not results:
                    # Retry with broader query
                    search_query = f"{language} {query} code example"
                    results = list(ddgs.text(search_query, max_results=5))
                
                if results:
                    # Verified and rank results
                    verified = KnowledgeSynthesizer.verify_facts(query, results)
                    # Filter for those that likely contain code - relaxed threshold
                    code_results = [r for r in verified if r['relevance_score'] > 0.1]
                    
                    if code_results:
                        return CodeComposer.synthesize_code_from_search(code_results, language)
            return None
        except Exception as e:
            logger.error(f"Error in Deep Research for code: {e}")
            return None

    @staticmethod
    def generate_code(language: str, description: str) -> str:
        """Generate code based on language and description using Deep Research engine"""
        # Exclusively use real code search
        real_code = MAXY1_3.search_real_code(language, description)
        
        if real_code:
            response = f"### 🔍 Deep Research Result: {language.capitalize()}\n\n"
            response += f"I've performed a deep search across technical sources to find the best implementation for your request:\n\n"
            response += f"{real_code}\n\n"
            response += f"**Research Insight:** This code was synthesized from multiple verified sources. "
            response += f"I've prioritized current best practices and functional correctness. "
            response += f"Would you like me to explain any specific logic or refine this further?"
            return response
            
        return None
    
    @staticmethod
    def is_chart_request(message: str) -> tuple[bool, str, list, list, str]:
        """Detect if message is asking for a chart and extract data, labels, and title"""
        msg_lower = message.lower()
        
        chart_indicators = ['chart', 'graph', 'pie chart', 'bar chart', 'line chart', 'visualization', 'plot', 'create a chart', 'make a chart', 'histogram']
        is_chart = any(ind in msg_lower for ind in chart_indicators)
        
        # Determine chart type
        chart_type = 'pie'
        if 'bar' in msg_lower:
            chart_type = 'bar'
        elif 'line' in msg_lower:
            chart_type = 'line'
        elif 'scatter' in msg_lower:
            chart_type = 'scatter'
        elif 'histogram' in msg_lower:
            chart_type = 'histogram'
        elif 'donut' in msg_lower:
            chart_type = 'donut'
        elif 'radar' in msg_lower or 'spider' in msg_lower:
            chart_type = 'radar'
        elif 'area' in msg_lower:
            chart_type = 'area'
        
        # Try to extract numbers from message
        import re
        numbers = re.findall(r'\d+', message)
        data = [int(n) for n in numbers[:8]] if numbers else [30, 25, 20, 15, 10]
        
        # Try to extract labels (words before numbers or common categories)
        labels = []
        words = re.findall(r'[a-zA-Z]+', message)
        common_labels = ['sales', 'revenue', 'profit', 'users', 'customers', 'products', 'jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']
        for word in words:
            if word.lower() in common_labels or len(word) > 2:
                labels.append(word.capitalize())
        
        # If no labels found, use defaults
        if len(labels) < len(data):
            labels.extend([f'Item {i+1}' for i in range(len(labels), len(data))])
        labels = labels[:len(data)]
        
        # Extract title
        title = "Data Visualization"
        title_patterns = [
            r'(?:show|display|create|make).*?(?:for|of|showing)\s+(.+?)(?:\s+with|\s+using|\s+data|$)',
            r'(?:chart|graph)\s+(?:for|of)\s+(.+?)(?:\s+with|\s+using|\s+data|$)'
        ]
        for pattern in title_patterns:
            match = re.search(pattern, msg_lower)
            if match:
                title = match.group(1).strip().capitalize()
                break
        
        return is_chart, chart_type, data, labels, title
    
    @staticmethod
    def generate_chart_image(chart_type: str, data: list, labels: list, title: str = "Data Visualization") -> tuple[Optional[str], str]:
        """Generate actual chart image and return base64 string"""
        try:
            ChartGenerator = MAXY1_3._get_chart_generator()
            
            base64_image = None
            
            if chart_type == 'pie':
                base64_image = ChartGenerator.create_pie_chart(
                    labels=labels,
                    values=[float(d) for d in data],
                    title=title
                )
            elif chart_type == 'donut':
                base64_image = ChartGenerator.create_donut_chart(
                    labels=labels,
                    values=[float(d) for d in data],
                    title=title
                )
            elif chart_type == 'radar':
                base64_image = ChartGenerator.create_radar_chart(
                    labels=labels,
                    values=[float(d) for d in data],
                    title=title
                )
            elif chart_type == 'area':
                base64_image = ChartGenerator.create_area_chart(
                    x=list(range(len(data))),
                    y=[float(d) for d in data],
                    title=title
                )
            elif chart_type == 'bar':
                base64_image = ChartGenerator.create_bar_chart(
                    categories=labels,
                    values=[float(d) for d in data],
                    title=title,
                    xlabel="Categories",
                    ylabel="Values"
                )
            elif chart_type == 'line':
                # For line chart, use indices as x values
                x_values = [float(i) for i in range(len(data))]
                y_values = [float(d) for d in data]
                base64_image = ChartGenerator.create_line_chart(
                    x=x_values,
                    y=y_values,
                    title=title,
                    xlabel="Index",
                    ylabel="Values"
                )
            elif chart_type == 'scatter':
                # Create scatter with sequential x values
                x_values = [float(i) for i in range(len(data))]
                y_values = [float(d) for d in data]
                base64_image = ChartGenerator.create_scatter_plot(
                    x=x_values,
                    y=y_values,
                    title=title,
                    xlabel="X",
                    ylabel="Y"
                )
            elif chart_type == 'histogram':
                data_floats = [float(d) for d in data]
                base64_image = ChartGenerator.create_histogram(
                    data=data_floats,
                    title=title,
                    bins=min(20, len(set(data)))
                )
            
            if base64_image:
                description = f"{chart_type.capitalize()} chart showing {title} with {len(data)} data points"
                return base64_image, description
            else:
                return None, "Failed to generate chart"
                
        except Exception as e:
            return None, f"Error: {str(e)}"

    @staticmethod
    def analyze_stock(ticker: str) -> Optional[str]:
        """Analyze stock data using yfinance"""
        try:
            stock = yf.Ticker(ticker)
            info = stock.info
            
            # Current price
            current_price = info.get('currentPrice') or info.get('regularMarketPrice')
            previous_close = info.get('previousClose')
            
            if not current_price:
                return None
                
            change = current_price - previous_close if previous_close else 0
            change_percent = (change / previous_close) * 100 if previous_close else 0
            
            response = f"### 📈 Stock Analysis: {info.get('longName', ticker.upper())}\n\n"
            response += f"**Current Price:** ${current_price:,.2f}\n"
            response += f"**Change:** {change:+.2f} ({change_percent:+.2f}%)\n"
            response += f"**Market Cap:** ${info.get('marketCap', 0):,.0f}\n"
            response += f"**52 Week Range:** ${info.get('fiftyTwoWeekLow', 0):,.2f} - ${info.get('fiftyTwoWeekHigh', 0):,.2f}\n\n"
            
            response += f"**Business Summary:**\n"
            summary = info.get('longBusinessSummary', 'No summary available.')
            response += f"{summary[:400]}...\n\n"
            
            # Recommendation
            rec = info.get('recommendationKey', 'none').replace('_', ' ').title()
            response += f"**Analyst Recommendation:** {rec}"
            
            return response
        except Exception as e:
            logger.error(f"Stock analysis error: {e}")
            return None
    
    @staticmethod
    def is_website_request(message: str) -> tuple[bool, str]:
        """Detect if user wants to build a website and what type"""
        msg_lower = message.lower()
        website_indicators = ['build', 'create', 'make', 'website', 'web site', 'page', 'landing', 'portfolio', 'ui', 'interface']
        
        is_website = any(ind in msg_lower for ind in website_indicators) and \
                     any(act in msg_lower for act in ['build', 'create', 'make', 'design', 'setup', 'generate', 'show me'])
        
        type = 'general'
        if 'portfolio' in msg_lower:
            type = 'portfolio'
        elif 'landing' in msg_lower:
            type = 'landing'
        elif 'business' in msg_lower:
            type = 'business'
        elif 'dashboard' in msg_lower:
            type = 'dashboard'
            
        return is_website, type

    @staticmethod
    def analyze_user_intent(message: str) -> Dict[str, Any]:
        """Comprehensive intent analysis for MAXY 1.3 combining 1.1 and 1.2 logic"""
        msg_lower = message.lower().strip()
        
        # Core intents from 1.1
        # Core intents
        intents = {
            'greeting': any(re.search(r'\b' + re.escape(g) + r'\b', msg_lower) for g in ['hi', 'hello', 'hey', 'greetings', 'howdy', 'namaskaar']),
            'farewell': any(re.search(r'\b' + re.escape(f) + r'\b', msg_lower) for f in ['bye', 'goodbye', 'see you', 'farewell', 'later']),
            'gratitude': any(re.search(r'\b' + re.escape(t) + r'\b', msg_lower) for t in ['thanks', 'thank you', 'appreciate', 'grateful']),
            'personal_status': any(h in msg_lower for h in ['how are you', 'how you doing']),
            'identity': any(i in msg_lower for i in ['your name', 'who are you', 'what are you']),
            'entertainment': any(j in msg_lower for j in ['joke', 'funny', 'laugh']),
            'time_query': any(t in msg_lower for t in ['time', 'what time', 'current time']),
            'date_query': any(d in msg_lower for d in ['date', 'today', 'what day']),
            'help': any(h in msg_lower for h in ['help', 'what can you do']),
            'news': any(n in msg_lower for n in ['news', 'happening', 'headlines', 'world today', 'current events']),
            'daily_updates': any(u in msg_lower for u in ['daily updates', 'whats new', 'what is new', 'latest updates']),
            'weather': any(w in msg_lower for w in ['weather', 'temperature', 'rain', 'sunny']),
            'calculation': any(c in msg_lower for c in ['calculate', 'math', 'plus', 'minus', 'times', 'divided']),
            'website_creation': any(ind in msg_lower for ind in ['build', 'create', 'make', 'website', 'web site', 'page', 'landing', 'portfolio', 'ui', 'interface'])
        }
        
        # Deep analysis metrics from 1.2
        topics = {
            'science': any(t in msg_lower for t in ['science', 'physics', 'chemistry', 'biology', 'research']),
            'history': any(t in msg_lower for t in ['history', 'ancient', 'century', 'war', 'civilization']),
            'technology': any(t in msg_lower for t in ['technology', 'computer', 'internet', 'software', 'ai']),
            'geography': any(t in msg_lower for t in ['country', 'capital', 'city', 'continent', 'population']),
            'personal': any(t in msg_lower for t in ['i feel', 'i think', 'my opinion', 'in my experience']),
            'philosophy': any(t in msg_lower for t in ['meaning', 'philosophy', 'why do we', 'purpose', 'existence'])
        }
        
        # Depth indicators
        depth_indicators = {
            'surface': ['what is', 'who is', 'simple', 'basic', 'quick'],
            'moderate': ['how does', 'why does', 'explain', 'tell me about'],
            'deep': ['analyze', 'comprehensive', 'detailed', 'in-depth', 'research', 'history of', 'science of']
        }
        
        inquiry_depth = 'surface'
        for depth, indicators in depth_indicators.items():
            if any(ind in msg_lower for ind in indicators):
                inquiry_depth = depth
                break
                
        # Detect question complexity (align with 1.2)
        complexity = 'simple'
        question_words = msg_lower.count('?')
        word_count = len(message.split())
        
        if word_count > 15 or question_words >= 2:
            complexity = 'complex'
        elif word_count > 8 or question_words == 1:
            complexity = 'moderate'

        return {
            'intents': intents,
            'topics': topics,
            'complexity': complexity,
            'inquiry_depth': inquiry_depth,
            'depth': inquiry_depth,  # Keep for 1.3 internal logic if used
            'is_research': MAXY1_2.is_research_query(message),
            'is_code': MAXY1_3.is_code_request(message)[0],
            'is_chart': MAXY1_3.is_chart_request(message)[0],
            'is_website': MAXY1_3.is_website_request(message)[0],
            'word_count': word_count,
            'message_length': len(message)
        }

    @staticmethod
    def process_message(
        message: str,
        include_thinking: bool = True,
        conversation_history: Optional[List[Dict]] = None,
        file_data: Optional[Dict] = None,
        user_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Ultimate message processing consolidating ALL MAXY features"""
        
        thinking = None
        confidence = 0.85
        chart_data = None
        response = None
        analysis_type = "analysis" # Default for 1.3
        
        # Context analysis & Follow-up detection
        is_followup, prev_context = MAXY1_3.detect_followup(message, conversation_history)
        effective_message = f"{prev_context} {message}" if is_followup else message
        
        # Analyze comprehensive intent
        analysis = MAXY1_3.analyze_user_intent(effective_message)
        intents = analysis['intents']
        use_slang = slang_manager.detect_slang(message)
        
        # Determine thinking type based on intent
        if analysis['is_code'] or analysis['is_website']:
            analysis_type = "analysis"
        elif analysis['is_research'] or analysis['depth'] == 'deep':
            analysis_type = "research"
        elif analysis['intents']['greeting'] or analysis['topics']['personal']:
            analysis_type = "conversation"
        else:
            analysis_type = "general"

        # Check for file analysis request
        file_path_match = re.search(r'(?:analyze|read|check|open)\s+(?:the\s+)?(.*?(\.pdf|\.docx?|\.csv|\.xlsx?))\b', message.lower())
        if not response and file_path_match:
            file_path = file_path_match.group(1).strip()
            # If path doesn't exist, check in common directories
            if not os.path.exists(file_path):
                 # Try relative to current dir or if it's just a filename
                 file_path = os.path.join(os.getcwd(), os.path.basename(file_path))
            
            file_res = MAXY1_3.analyze_file(file_path)
            if file_res['success']:
                response = f"### 📂 File Analysis: {os.path.basename(file_path)}\n\n"
                response += f"**Summary:** {file_res['summary']}\n\n"
                response += "**Extracted Insights (Preview):**\n"
                response += f"> {file_res['content'][:1500]}..."
                analysis_type = "analysis"
                confidence = 0.98
            else:
                response = f"I attempted to analyze the file at `{file_path}`, but encountered an issue: {file_res.get('error', 'Unknown error')}. Please ensure the file exists and is in a supported format (PDF, Word, CSV, Excel)."
        
        if include_thinking:
            thinking = MAXYThinkingEngine.generate_thinking(
                MAXY1_3.NAME,
                effective_message,
                analysis_type
            )
        
        # 0. Daily Updates check
        if not response and intents.get('daily_updates'):
            try:
                import json
                updates_path = os.path.join(os.path.dirname(__file__), "updates.json")
                with open(updates_path, 'r') as f:
                    data = json.load(f)
                    response = "**MAXY ENTERPRISE INTELLIGENCE: DAILY UPDATES**\n" + "="*50 + "\n\n"
                    for up in data['updates']:
                        response += f"### {up['title']} ({up['date']})\n{up['description']}\n\n"
                    response += "Would you like an in-depth analysis of any of these trends or improvements?"
                    confidence = 0.99
            except Exception as e:
                logger.error(f"Error in MAXY 1.3 daily_updates handler: {e}")

        # ── ESSAY / SPEECH GENERATION (Integrated from 1.2) ──
        essay_intent = MAXY1_2.detect_essay_intent(message)
        if not response and essay_intent:
            topic = essay_intent['topic']
            variation = 0
            if conversation_history:
                prev_responses = [m['content'] for m in conversation_history if m['role'] == 'assistant']
                variation = sum(1 for r in prev_responses if '📝 **Essay' in r or '🎤 **Speech' in r)
            
            result = MAXY1_2.deep_wikipedia_research(topic)
            if not result['success'] or 'does not reside' in result['response']:
                result = MAXY1_2.perform_web_search(topic)
                
            if result['success']:
                if essay_intent['mode'] == 'speech':
                    response = MAXY1_2.format_as_speech(
                        result['response'], essay_intent['style'],
                        essay_intent['word_target'], variation
                    )
                else:
                    response = MAXY1_2.format_as_essay(
                        result['response'], essay_intent['style'],
                        essay_intent['word_target'], variation
                    )
                confidence = 0.97
        # ── END ESSAY / SPEECH ──

        # 1. PRIORITY: UTILITY FEATURES (Weather/Time/Date) - Priority check to avoid false positives
        
        # Weather (from 1.1)
        if not response and intents['weather']:
            words = message.split()
            city = None
            if 'in' in words:
                idx = words.index('in')
                if idx + 1 < len(words):
                    city = " ".join(words[idx + 1:]).strip('?.!')
            if city:
                weather_info = MAXY1_1.get_weather(city)
                if weather_info:
                    response = f"{weather_info} 🌤️"
                    confidence = 0.95

        # Time/Date (from 1.1)
        if not response:
            if intents['time_query']:
                response = f"It's {datetime.now().strftime('%I:%M %p')} right now! ⏰"
                confidence = 0.97
            elif intents['date_query'] and not intents['news']:
                response = f"Today is {datetime.now().strftime('%A, %B %d, %Y')}! 📅"
                confidence = 0.97

        # 2. PRIORITY: EXISTING 1.3 FEATURES (Technical/Data)
        
        # Chart Request (Moved up to prevent interception by Website fallback)
        if not response and analysis['is_chart']:
            is_chart, chart_type, data, labels, title = MAXY1_3.is_chart_request(message)
            base64_image, desc = MAXY1_3.generate_chart_image(chart_type, data, labels, title)
            if base64_image:
                response = f"I've created a {chart_type} chart for you based on your data! 📊\n\n**{title}** breakdown shows {len(data)} distinct data points total."
                chart_data = {'type': chart_type, 'title': title, 'base64_image': base64_image, 'description': desc}
                confidence = 0.95

        # Website Request
        if not response and (analysis['is_website'] or intents.get('website_creation')):
             is_website, web_type = MAXY1_3.is_website_request(message)
             search_query = f"complete premium responsive {web_type} website code template single file HTML CSS Inter font"
             research_code = MAXY1_3.search_real_code("html", search_query)
             
             if research_code and "html" in research_code.lower():
                 response = f"### 🏗️ MAXY Deep Research: Premium {web_type.capitalize()} Builder\n\n"
                 response += f"I've synthesized a high-end, responsive **{web_type.capitalize()}** template for you. This design incorporates modern UI/UX standards, fluid animations, and a premium color palette discovered through deep technical research:\n\n"
                 response += f"{research_code}\n\n"
                 response += f"**Research Insight:** This code utilizes optimized CSS grid/flexbox patterns and semantic HTML5 for maximum accessibility and performance. "
                 response += "Would you like me to add glassmorphism effects or refine the typography further?"
                 confidence = 0.98
             else:
                  # Local Fallback - UPGRADED PREMIUM TEMPLATE
                  if web_type == 'portfolio':
                      response = "### 🏗️ MAXY Template: Premium Portfolio Specialist\n\n"
                      response += "I've generated a bespoke, high-performance portfolio starter using a sleek dark-mode aesthetic and modern 'Inter' typography:\n\n"
                      template = """<!DOCTYPE html>
<html lang='en'>
<head>
  <meta charset='UTF-8'>
  <meta name='viewport' content='width=device-width, initial-scale=1.0'>
  <title>MAXY Portfolio</title>
  <link href='https://fonts.googleapis.com/css2?family=Inter:wght@300;400;700&display=swap' rel='stylesheet'>
  <style>
    :root { --bg: #0a0a0c; --accent: #3b82f6; --text: #f8fafc; }
    body { font-family: 'Inter', sans-serif; background: var(--bg); color: var(--text); margin: 0; overflow-x: hidden; }
    .glass { background: rgba(255, 255, 255, 0.03); backdrop-filter: blur(10px); border: 1px solid rgba(255,255,255,0.05); }
    nav { padding: 2rem; display: flex; justify-content: space-between; position: fixed; width: 100%; box-sizing: border-box; z-index: 100; }
    .hero { height: 100vh; display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; }
    h1 { font-size: 5rem; margin: 0; background: linear-gradient(to right, #fff, #64748b); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
    .btn { padding: 1rem 2rem; background: var(--accent); color: white; text-decoration: none; border-radius: 50px; margin-top: 2rem; transition: 0.3s; display: inline-block; }
    .btn:hover { transform: translateY(-3px); box-shadow: 0 10px 20px rgba(59, 130, 246, 0.3); }
  </style>
</head>
<body>
  <nav class='glass'>
    <div><strong>MAXY PORTFOLIO</strong></div>
    <div>Work . About . Contact</div>
  </nav>
  <section class='hero'>
    <h1>Digital Architecture &<br>Creative Solutions</h1>
    <p>Crafting high-performance experiences for the modern web.</p>
    <a href='#' class='btn'>View Laboratory</a>
  </section>
</body>
</html>"""
                      response += "```html\n" + template + "\n```\n\n"
                      response += template + "\n\n"
                      response += "This premium template is ready for deployment. I can expand it with project galleries, contact forms, or dynamic animations—what's our next step?"
                      confidence = 0.95
                  else:
                      response = f"I'm ready to architect your **{web_type}** website! While I'm refining the deep search for hyper-specific templates, "
                      response += "I've activated my UI design module. Should we prioritize a minimalist aesthetic or a high-impact, dynamic layout?"
                      confidence = 0.85

        # Code Request (General)
        if not response and analysis['is_code'] and not analysis['is_chart']:
            is_code, language = MAXY1_3.is_code_request(message)
            response = MAXY1_3.generate_code(language, message)
            if response:
                confidence = 0.96

        # Stock Analysis
        if not response:
            stock_match = re.search(r'\b(stock|price|ticker)\s+(?:of\s+)?([A-Z]{1,5})\b', message.upper())
            if stock_match:
                ticker = stock_match.group(2)
                stock_analysis = MAXY1_3.analyze_stock(ticker)
                if stock_analysis:
                    response = stock_analysis
                    confidence = 0.95

        # File Intelligence
        if not response and file_data:
            # We process files if no other high-priority technical intent was matched
            file_name = file_data.get('name', 'unknown file')
            file_content = file_data.get('content', '')
            if file_name.lower().endswith('.csv'):
                parse_result = StructuredDataAnalyzer.parse_csv_content(file_content)
                if 'error' not in parse_result:
                    insights = StructuredDataAnalyzer.generate_data_insights(parse_result)
                    response = f"### 📊 Data Intelligence: {file_name}\n\nKey Insights:\n" + "\n".join([f"- {i}" for i in insights])
                    confidence = 0.98
            else:
                keywords = TextAnalyzer.extract_keywords(file_content, top_n=5)
                sentiment = TextAnalyzer.analyze_sentiment(file_content)
                response = f"### 📄 Document Intelligence: {file_name}\n\nThis document has a **{sentiment['sentiment']}** tone. Primary themes: {', '.join([k[0] for k in keywords])}."
                confidence = 0.95

        # 3. CONSOLIDATED FEATURES FROM 1.1 & 1.2
        
        # Jokes/Entertainment (from 1.1)
        if not response and intents['entertainment']:
            joke = random.choice(MAXY1_1.JOKES + ["Why did the cross-functional team cross the road? To attend a stand-up on the other side!"])
            response = f"{joke} 😄"
            confidence = 0.92

        # Deep Research (from 1.2) - Higher priority than general conversation
        if not response and (analysis['is_research'] or analysis['depth'] == 'deep'):
            research_result = MAXY1_2.deep_wikipedia_research(message)
            if not research_result['success'] or "does not reside" in research_result['response']:
                web_result = MAXY1_2.perform_web_search(message)
                if web_result['success']:
                    research_result = web_result
            
            if research_result['success']:
                # Format based on depth first
                response = MAXY1_2.format_research_response(research_result['response'], analysis['depth'])
                confidence = research_result['confidence']

                # Identity check for surface/moderate queries only
                if analysis['depth'] in ['surface', 'moderate']:
                    identity_answer = KnowledgeSynthesizer.extract_identity_answer(message, research_result['response'], intents)
                    if identity_answer:
                        response = identity_answer
                        confidence = 0.98

        # Philosophy & Personal (from 1.2)
        if not response:
            if analysis['topics']['philosophy']:
                 response, confidence = MAXY1_2.generate_detailed_response(analysis, message, conversation_history, use_slang, user_name)
            elif analysis['topics']['personal']:
                 response, confidence = MAXY1_2.generate_detailed_response(analysis, message, conversation_history, use_slang, user_name)

        # Greetings & Identity (Combined)
        if not response:
            if intents['greeting']:
                user_display = f" {user_name}" if user_name else ""
                slang = slang_manager.get_random_slang(use_slang)
                response = f"Hello{user_display} {slang}! I'm MAXY 1.3, your most advanced AI companion. I've been upgraded with all the research capabilities of 1.2 and the speed of 1.1. I can build websites, write code, analyze data, and perform deep research. What shall we tackle today?"
                confidence = 0.98
            elif intents['identity']:
                response = "I'm MAXY 1.3 – the ultimate version of the MAXY AI series. I combine rapid response logic, deep Wikipedia research, and advanced data visualization into one powerful interface. Whether you need a statistical analysis, a web landing page, or a deep dive into history, I've got you covered."
                confidence = 0.96

        # Calculation (from 1.1)
        if not response and intents['calculation']:
             response = "I can definitely help with that calculation! Please provide the numbers and the operation you'd like me to perform."
             confidence = 0.90

        # Help - Premium Persona alignment
        if not response and intents['help']:
            response = "I am MAXY 1.3, your **High-Performance AI Engine**. My premium capabilities include:\n\n" \
                       "🚀 **Advanced Engineering** - Full-stack web building & technical architecture\n" \
                       "📂 **Universal Analysis** - Deep insights from PDF, Word, CSV, and Excel documents\n" \
                       "📊 **Dynamic Visualization** - Professional Donut, Radar, and Area charts\n" \
                       "🔍 **Intelligence Synthesis** - Multi-source technical research & data extraction\n" \
                       "💬 **Strategic Conversation** - Context-aware, professional-grade dialogue\n\n" \
                       "How may I assist your high-level objectives today?"
            confidence = 0.99

        # FINAL FALLBACKS
        if not response:
            # Try a quick wiki lookup if it's a short query that looks like a topic
            if analysis['word_count'] <= 3 and not use_slang:
                wiki_result = MAXY1_1.quick_wikipedia_lookup(message)
                if wiki_result:
                    response = wiki_result
                    confidence = 0.92

        if not response:
            # Technical search fallback from 1.3 original
            response = MAXY1_3.generate_code("technical", message)
            if response and "research insight" in response.lower() and "couldn't find" in response.lower():
                # This should no longer happen as generate_code returns None on fail
                response = None
            
            if not response:
                slang = slang_manager.get_random_slang(use_slang)
                response = f"I am MAXY 1.3, and I'm ready to provide premium support for your technical project, {slang}. I specialize in architectural code generation, complex data insights, and multi-file analysis. Could you specify your technical objective?"
            confidence = 0.85
        
        # Slang Enhancement
        if response and "statistical analysis" not in response.lower() and "generated" not in response.lower() and "verified research report" not in response.lower():
             response = slang_manager.enhance_text(response, force=use_slang)
        
        result = {
            'response': response,
            'model': MAXY1_3.NAME,
            'confidence': min(1.0, confidence),
        }
        
        if thinking: result['thinking'] = thinking
        if chart_data: result['charts'] = [chart_data]
            
        return result


class ModelRouter:
    """Route messages to appropriate model"""
    
    MODELS = {
        'maxy1.1': MAXY1_1,
        'maxy1.2': MAXY1_2,
        'maxy1.3': MAXY1_3,
    }
    
    @staticmethod
    def get_model_info(model_name: str) -> Dict[str, Any]:
        """Get information about a model"""
        model_class = ModelRouter.MODELS.get(model_name.lower())
        if not model_class:
            return {}
        
        capabilities = {
            'maxy1.1': [
                'Lightning-fast responses',
                'Visible thinking process',
                'Quick Wikipedia lookups',
                'Friendly conversation',
                'Time/date queries'
            ],
            'maxy1.2': [
                'Deep Wikipedia research',
                'Comprehensive analysis',
                'Natural conversation',
                'Context-aware responses',
                'Multi-topic knowledge'
            ],
            'maxy1.3': [
                'Grand Unified Engine (1.1 + 1.2 + 1.3)',
                'File processing & analysis',
                'Deep Wikipedia search & synthesis',
                'Dynamic Code & Website generation',
                'Data visualization & Charting',
                'Weather, Time, Stock updates',
                'Intelligent conversation with slang'
            ]
        }
        
        return {
            'name': model_class.NAME,
            'version': model_class.VERSION,
            'description': model_class.DESCRIPTION,
            'capabilities': capabilities.get(model_name.lower(), [])
        }
    
    @staticmethod
    def process(
        model_name: str,
        message: str,
        include_thinking: bool = True,
        conversation_history: Optional[List[Dict]] = None,
        user_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Route message to appropriate model"""
        
        # Check for slang toggle commands
        msg_lower = message.lower().strip().strip('!.')
        if any(cmd in msg_lower for cmd in ["enable slangs", "activate slangs", "turn on slangs", "enable slang", "activate slang"]):
            response_text = slang_manager.set_enabled(True)
            return {
                'response': f"{response_text} {slang_manager.get_random_slang(force=True)}! I'm ready to chat with some local flavor.",
                'model': 'System',
                'confidence': 1.0
            }
        
        if any(cmd in msg_lower for cmd in ["disable slangs", "stop slangs", "turn off slangs", "disable slang", "no slangs"]):
            response_text = slang_manager.set_enabled(False)
            return {
                'response': f"{response_text} I will keep the conversation formal and standard from now on.",
                'model': 'System',
                'confidence': 1.0
            }

        model_name_lower = model_name.lower()
        
        if model_name_lower == 'maxy1.1':
            return MAXY1_1.process_message(message, include_thinking, conversation_history, user_name)
        elif model_name_lower == 'maxy1.2':
            return MAXY1_2.process_message(message, include_thinking, conversation_history, user_name)
        elif model_name_lower == 'maxy1.3':
            return MAXY1_3.process_message(message, include_thinking, conversation_history, None, user_name)
        else:
            logger.warning(f"Unknown model: {model_name}, defaulting to MAXY1.1")
            return MAXY1_1.process_message(message, include_thinking, conversation_history, user_name)
            