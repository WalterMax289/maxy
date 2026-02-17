"""
Advanced file processing for various document types
Supports PDF, Word, images, CSV, JSON, and text files
"""

import base64
import io
import logging
from typing import Dict, Any, Optional, Tuple
import json

logger = logging.getLogger(__name__)

# Optional imports with availability tracking
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("PIL not available - image processing disabled")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available - PDF processing disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word document processing disabled")

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False


class FileProcessor:
    """Process and analyze various file types"""
    
    @staticmethod
    def process_image(base64_content: str) -> Dict[str, Any]:
        """Analyze image and extract detailed information"""
        try:
            if not PIL_AVAILABLE:
                return {
                    'success': False,
                    'error': 'PIL (Pillow) not available',
                    'suggestion': 'Install with: pip install Pillow'
                }
            
            # Decode base64
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]
            
            image_data = base64.b64decode(base64_content)
            image = Image.open(io.BytesIO(image_data))
            
            # Extract image information
            width, height = image.size
            format_type = image.format or "Unknown"
            mode = image.mode
            aspect_ratio = width / height if height > 0 else 0
            file_size = len(image_data)
            
            # Analyze color information
            color_info = FileProcessor._analyze_image_colors(image, mode)
            
            # Categorize image
            size_category = FileProcessor._categorize_image_size(width, height)
            
            # Get EXIF data if available
            exif_data = FileProcessor._extract_exif_data(image)
            
            analysis = (
                f"IMAGE ANALYSIS REPORT\n"
                f"{'='*50}\n\n"
                f"DIMENSIONS & FORMAT\n"
                f"• Dimensions: {width} × {height} pixels\n"
                f"• Aspect Ratio: {aspect_ratio:.2f}:1 "
                f"({'Landscape' if aspect_ratio > 1.2 else 'Portrait' if aspect_ratio < 0.8 else 'Square'})\n"
                f"• Format: {format_type}\n"
                f"• Color Mode: {mode}\n"
                f"• Size: {file_size / 1024:.1f} KB\n"
                f"• Category: {size_category}\n\n"
                f"COLOR ANALYSIS\n"
                f"{color_info}\n"
            )
            
            if exif_data:
                analysis += f"METADATA\n{exif_data}\n"
            
            return {
                'success': True,
                'analysis': analysis,
                'metadata': {
                    'width': width,
                    'height': height,
                    'format': format_type,
                    'mode': mode,
                    'size_kb': file_size / 1024,
                    'aspect_ratio': aspect_ratio,
                    'category': size_category
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            return {
                'success': False,
                'error': f'Image processing failed: {str(e)}',
                'suggestion': 'Ensure file is a valid image format (PNG, JPG, GIF, etc.)'
            }
    
    @staticmethod
    def process_pdf(base64_content: str) -> Dict[str, Any]:
        """Extract and analyze PDF document"""
        try:
            if not PDF_AVAILABLE:
                return {
                    'success': False,
                    'error': 'PyPDF2 not available',
                    'suggestion': 'Install with: pip install PyPDF2'
                }
            
            # Decode base64
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]
            
            pdf_data = base64.b64decode(base64_content)
            pdf_file = io.BytesIO(pdf_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            num_pages = len(pdf_reader.pages)
            text_content = ""
            
            # Extract text from first 10 pages
            for i, page in enumerate(pdf_reader.pages[:10]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {i+1} ---\n{page_text}"
                except:
                    continue
            
            # Extract metadata
            metadata = pdf_reader.metadata
            title = metadata.get('/Title', 'Unknown') if metadata else 'Unknown'
            author = metadata.get('/Author', 'Unknown') if metadata else 'Unknown'
            subject = metadata.get('/Subject', 'Unknown') if metadata else 'Unknown'
            creator = metadata.get('/Creator', 'Unknown') if metadata else 'Unknown'
            
            word_count = len(text_content.split())
            
            analysis = (
                f"PDF DOCUMENT ANALYSIS REPORT\n"
                f"{'='*50}\n\n"
                f"DOCUMENT INFORMATION\n"
                f"• Title: {title}\n"
                f"• Author: {author}\n"
                f"• Subject: {subject}\n"
                f"• Creator: {creator}\n\n"
                f"STATISTICS\n"
                f"• Total Pages: {num_pages}\n"
                f"• Extracted Pages: {min(10, num_pages)}\n"
                f"• Approximate Word Count: {word_count}\n"
                f"• Average Words per Page: {word_count // max(num_pages, 1)}\n\n"
                f"CONTENT PREVIEW (First 1000 characters)\n"
                f"{'-'*50}\n"
                f"{text_content[:1000]}{'...' if len(text_content) > 1000 else ''}\n"
            )
            
            return {
                'success': True,
                'analysis': analysis,
                'metadata': {
                    'pages': num_pages,
                    'word_count': word_count,
                    'title': title,
                    'author': author,
                    'subject': subject
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return {
                'success': False,
                'error': f'PDF processing failed: {str(e)}',
                'suggestion': 'Ensure file is a valid PDF document'
            }
    
    @staticmethod
    def process_word_document(base64_content: str) -> Dict[str, Any]:
        """Extract and analyze Word document"""
        try:
            if not DOCX_AVAILABLE:
                return {
                    'success': False,
                    'error': 'python-docx not available',
                    'suggestion': 'Install with: pip install python-docx'
                }
            
            # Decode base64
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]
            
            doc_data = base64.b64decode(base64_content)
            doc_file = io.BytesIO(doc_data)
            doc = docx.Document(doc_file)
            
            # Extract text
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            text_content = '\n'.join(full_text)
            word_count = len(text_content.split())
            para_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Extract tables if present
            table_count = len(doc.tables)
            table_info = ""
            if table_count > 0:
                table_info = f"\n• Tables: {table_count}"
            
            analysis = (
                f"WORD DOCUMENT ANALYSIS REPORT\n"
                f"{'='*50}\n\n"
                f"DOCUMENT STATISTICS\n"
                f"• Paragraphs: {para_count}\n"
                f"• Approximate Word Count: {word_count}\n"
                f"• Character Count: {len(text_content)}{table_info}\n\n"
                f"CONTENT PREVIEW (First 1500 characters)\n"
                f"{'-'*50}\n"
                f"{text_content[:1500]}{'...' if len(text_content) > 1500 else ''}\n"
            )
            
            return {
                'success': True,
                'analysis': analysis,
                'metadata': {
                    'paragraphs': para_count,
                    'word_count': word_count,
                    'character_count': len(text_content),
                    'tables': table_count
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            return {
                'success': False,
                'error': f'Word document processing failed: {str(e)}',
                'suggestion': 'Ensure file is a valid .docx or .doc document'
            }
    
    @staticmethod
    def process_text_file(base64_content: str, filename: str) -> Dict[str, Any]:
        """Process and analyze text files (TXT, CSV, JSON, etc.)"""
        try:
            # Decode base64
            if ',' in base64_content:
                base64_content = base64_content.split(',')[1]
            
            text = base64.b64decode(base64_content).decode('utf-8', errors='ignore')
            
            word_count = len(text.split())
            line_count = len(text.split('\n'))
            char_count = len(text)
            
            # Determine file type and parse accordingly
            analysis = FileProcessor._parse_by_extension(filename, text)
            
            base_analysis = (
                f"TEXT FILE ANALYSIS REPORT\n"
                f"{'='*50}\n\n"
                f"FILE INFORMATION\n"
                f"• Filename: {filename}\n"
                f"• Characters: {char_count}\n"
                f"• Words: {word_count}\n"
                f"• Lines: {line_count}\n\n"
            )
            
            if analysis:
                base_analysis += analysis + "\n\n"
            
            base_analysis += (
                f"CONTENT PREVIEW (First 1500 characters)\n"
                f"{'-'*50}\n"
                f"{text[:1500]}{'...' if len(text) > 1500 else ''}\n"
            )
            
            return {
                'success': True,
                'analysis': base_analysis,
                'metadata': {
                    'filename': filename,
                    'character_count': char_count,
                    'word_count': word_count,
                    'line_count': line_count
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            return {
                'success': False,
                'error': f'Text file processing failed: {str(e)}',
                'suggestion': 'Ensure file is a valid text file with UTF-8 encoding'
            }
    
    @staticmethod
    def _analyze_image_colors(image: Image.Image, mode: str) -> str:
        """Analyze image color information"""
        try:
            if mode in ['RGB', 'RGBA']:
                small = image.resize((50, 50))
                pixels = list(small.getdata())
                
                if pixels:
                    if mode == 'RGB':
                        avg_r = sum(p[0] for p in pixels) // len(pixels)
                        avg_g = sum(p[1] for p in pixels) // len(pixels)
                        avg_b = sum(p[2] for p in pixels) // len(pixels)
                        avg_color = (avg_r, avg_g, avg_b)
                        return f"• Average Color (RGB): {avg_color}\n• Color Profile: RGB"
                    else:
                        avg_r = sum(p[0] for p in pixels) // len(pixels)
                        avg_g = sum(p[1] for p in pixels) // len(pixels)
                        avg_b = sum(p[2] for p in pixels) // len(pixels)
                        avg_a = sum(p[3] for p in pixels) // len(pixels)
                        avg_color = (avg_r, avg_g, avg_b, avg_a)
                        return f"• Average Color (RGBA): {avg_color}\n• Color Profile: RGBA with transparency"
            else:
                return f"• Color Mode: {mode}\n• Color Analysis: Not applicable for {mode} images"
        except:
            return f"• Color Analysis: Could not analyze ({mode})"
    
    @staticmethod
    def _categorize_image_size(width: int, height: int) -> str:
        """Categorize image by size"""
        max_dim = max(width, height)
        
        if max_dim < 100:
            return "Icon (< 100px)"
        elif max_dim < 300:
            return "Thumbnail (100-300px)"
        elif max_dim < 800:
            return "Small (300-800px)"
        elif max_dim < 1600:
            return "Medium (800-1600px)"
        elif max_dim < 3840:
            return "Large (1600-3840px)"
        else:
            return "Extra Large (> 3840px)"
    
    @staticmethod
    def _extract_exif_data(image: Image.Image) -> str:
        """Extract EXIF metadata from image"""
        try:
            exif_data = image._getexif()
            if exif_data:
                info_lines = []
                for tag_id, value in exif_data.items():
                    tag_name = Image.Exif.TAGS.get(tag_id, tag_id)
                    if tag_name and isinstance(value, str):
                        info_lines.append(f"• {tag_name}: {value[:50]}")
                if info_lines:
                    return '\n'.join(info_lines[:5])
            return ""
        except:
            return ""
    
    @staticmethod
    def _parse_by_extension(filename: str, content: str) -> str:
        """Parse file based on extension"""
        lower_name = filename.lower()
        
        if lower_name.endswith('.json'):
            try:
                data = json.loads(content)
                return (
                    f"JSON FILE ANALYSIS\n"
                    f"• Keys: {len(data) if isinstance(data, dict) else 'N/A'}\n"
                    f"• Type: {'Object' if isinstance(data, dict) else 'Array' if isinstance(data, list) else 'Value'}"
                )
            except:
                return "JSON FILE: Could not parse as valid JSON"
        
        elif lower_name.endswith('.csv'):
            lines = content.split('\n')
            return (
                f"CSV FILE ANALYSIS\n"
                f"• Rows: {len(lines)}\n"
                f"• Columns: {len(lines[0].split(',')) if lines else 0}\n"
                f"• Header: {lines[0] if lines else 'N/A'}"
            )
        
        elif lower_name.endswith(('.py', '.js', '.cpp', '.c', '.java')):
            lines = content.split('\n')
            code_lines = [l for l in lines if l.strip() and not l.strip().startswith('#')]
            return (
                f"CODE FILE ANALYSIS\n"
                f"• Total Lines: {len(lines)}\n"
                f"• Code Lines: {len(code_lines)}\n"
                f"• Language: {lower_name.split('.')[-1].upper()}"
            )
        
        return ""
    
    @staticmethod
    def detect_file_type(filename: str, mime_type: str) -> str:
        """Detect file type from name and MIME type"""
        lower_name = filename.lower()
        lower_mime = mime_type.lower()
        
        if lower_mime.startswith('image/') or any(lower_name.endswith(f'.{ext}') 
                                                   for ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']):
            return 'image'
        elif lower_mime == 'application/pdf' or lower_name.endswith('.pdf'):
            return 'pdf'
        elif 'word' in lower_mime or 'document' in lower_mime or lower_name.endswith(('.docx', '.doc')):
            return 'document'
        elif lower_mime.startswith('text/') or any(lower_name.endswith(f'.{ext}') 
                                                    for ext in ['txt', 'md', 'html', 'css', 'xml', 'json']):
            return 'text'
        elif lower_name.endswith(('.csv', '.xlsx', '.xls')) or 'sheet' in lower_mime:
            return 'data'
        else:
            return 'unknown'
